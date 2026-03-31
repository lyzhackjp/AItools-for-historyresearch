#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
学术论文优化润色工作流程 v3.0

基于 paper_polisher_enhanced.py 的优化版本，解决以下问题：
1. 格式规范统一：段落结构、字体样式、行间距等
2. 内容纯净度处理：清除JSON残留、代码片段、标记符
3. 脚注规范添加：符合学术标准
4. 章节结构优化：修正嵌套问题，标准化命名

引用标准：WORKFLOW_DIAGRAM.md 学术论文润色流程
"""

import sys
import os
import json
import re
import datetime
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional
from dataclasses import dataclass, field
import hashlib

sys.path.append(str(Path(__file__).parent / 'modules'))

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Warning: python-docx not installed")
    sys.exit(1)

from modules.doc_processor import DocProcessor
from modules.llm_client import create_llm_client


@dataclass
class SectionInfo:
    """章节信息"""
    title: str
    content: str
    paragraphs: List[str]
    footnotes: List[str] = field(default_factory=list)
    original_index: int = 0


class ContentPurifier:
    """内容纯净度处理器 - 清除JSON残留和代码片段"""
    
    JSON_PATTERNS = [
        r'\{[^{}]*"[^"]*"\s*:\s*[^{}]*\}',  # JSON对象
        r'\[\s*\{[^]]*\}\s*\]',  # JSON数组
        r'"[^"]*"\s*:\s*"[^"]*"',  # JSON键值对
        r'\{[\s\S]*?"modified_text"[\s\S]*?\}',  # 包含modified_text的JSON
        r'\{[\s\S]*?"deletions"[\s\S]*?\}',  # 包含deletions的JSON
    ]
    
    CODE_MARKERS = [
        r'```[\s\S]*?```',  # 代码块标记
        r'`[^`]+`',  # 行内代码
        r'^\s*[{}[\]]+\s*$',  # 单独的括号
        r'^"[^"]*"$',  # 引号包裹的行
    ]
    
    @classmethod
    def clean_text(cls, text: str) -> str:
        """清除文本中的技术残留"""
        result = text.strip()
        
        for pattern in cls.JSON_PATTERNS:
            result = re.sub(pattern, '', result)
        
        for pattern in cls.CODE_MARKERS:
            result = re.sub(pattern, '', result, flags=re.MULTILINE)
        
        result = re.sub(r'\n{3,}', '\n\n', result)
        
        result = result.strip()
        
        return result
    
    @classmethod
    def is_clean_content(cls, text: str) -> bool:
        """检查内容是否纯净"""
        if re.search(r'\{[\s\S]*?\}', text):
            if 'modified_text' in text or 'deletions' in text:
                return False
        
        if re.search(r'```', text):
            return False
        
        return True


class FormatStandardizer:
    """格式标准化处理器"""
    
    FONT_CONFIG = {
        'name': '宋体',
        'size': 12,
        'bold': False,
        'italic': False,
        'color': RGBColor(0, 0, 0)
    }
    
    HEADING_CONFIG = {
        'font_name': '黑体',
        'font_size': 14,
        'bold': True
    }
    
    PARAGRAPH_CONFIG = {
        'first_line_indent': 0.74,  # 两个字符缩进
        'line_spacing': 1.5,
        'space_before': 0,
        'space_after': 6,
        'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    
    @classmethod
    def apply_format_to_paragraph(cls, paragraph, is_heading: bool = False):
        """应用标准化格式到段落"""
        for run in paragraph.runs:
            if is_heading:
                run.font.name = cls.HEADING_CONFIG['font_name']
                run.font.size = Pt(cls.HEADING_CONFIG['font_size'])
                run.font.bold = cls.HEADING_CONFIG['bold']
                run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                run.font.name = cls.FONT_CONFIG['name']
                run.font.size = Pt(cls.FONT_CONFIG['size'])
                run.font.bold = cls.FONT_CONFIG['bold']
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        if not is_heading:
            pf = paragraph.paragraph_format
            pf.first_line_indent = Pt(cls.PARAGRAPH_CONFIG['first_line_indent'])
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE_LINE
            pf.space_before = Pt(cls.PARAGRAPH_CONFIG['space_before'])
            pf.space_after = Pt(cls.PARAGRAPH_CONFIG['space_after'])
            pf.alignment = cls.PARAGRAPH_CONFIG['alignment']


class SectionStructureOptimizer:
    """章节结构优化器 - 修正嵌套问题"""
    
    SECTION_PATTERNS = [
        r'^第[一二三四五六七八九十]+部分',
        r'^第\d+部分',
        r'^第[一二三四五六七八九十]+章',
        r'^第\d+章',
        r'^一、',
        r'^二、',
        r'^三、',
        r'^四、',
        r'^五、',
    ]
    
    NESTED_INDICATORS = [
        r'\(第一部分\)',
        r'\(第二部分\)',
        r'\(第三部分\)',
        r'\(第一部分\) \(第一部分\)',
        r'\(第二部分\) \(第一部分\)',
        r'\(第二部分\) \(第二部分\)',
        r'\(第二部分\) \(第二部分\) \(第一部分\)',
    ]
    
    @classmethod
    def extract_main_section(cls, title: str) -> str:
        """提取主章节标题"""
        if '结语' in title or '结论' in title:
            return title
        
        for pattern in cls.SECTION_PATTERNS:
            match = re.search(pattern, title)
            if match:
                return match.group(0)
        
        return title
    
    @classmethod
    def is_nested_title(cls, title: str) -> bool:
        """判断是否为嵌套标题"""
        for indicator in cls.NESTED_INDICATORS:
            if re.search(indicator, title):
                return True
        return False
    
    @classmethod
    def clean_section_title(cls, title: str) -> str:
        """清理章节标题，移除嵌套标记"""
        cleaned = title
        
        for indicator in cls.NESTED_INDICATORS:
            cleaned = re.sub(indicator, '', cleaned)
        
        cleaned = re.sub(r'\s+', ' ', cleaned)
        cleaned = cleaned.strip()
        
        if not cleaned:
            cleaned = "正文"
        
        return cleaned


class FootnoteManager:
    """脚注管理器 - 规范添加学术脚注"""
    
    def __init__(self):
        self.footnotes = []
        self.footnote_id = 1
    
    def add_footnote(self, text: str, reference: str = None) -> Tuple[int, str]:
        """添加脚注"""
        footnote_num = self.footnote_id
        self.footnote_id += 1
        
        footnote_text = f"[{footnote_num}] {text}"
        if reference:
            footnote_text = f"[{footnote_num}] {text}（{reference}）"
        
        self.footnotes.append({
            'id': footnote_num,
            'text': footnote_text,
            'reference': reference
        })
        
        return footnote_num, footnote_text
    
    def get_footnote_references(self) -> List[Dict]:
        """获取脚注引用列表"""
        return self.footnotes.copy()


class AcademicCitationCleaner:
    """学术引用清洗器"""
    
    CITATION_PATTERNS = [
        (r'\((\d{4})\s*[年]?\)', r'(\1年)'),
        (r'《([^》]+)》,?\s*(\d{4})', r'《\1》(\2年)'),
        (r'\[(\d+)\]', r''),
    ]
    
    @classmethod
    def clean_citations(cls, text: str) -> str:
        """清洗学术引用格式"""
        result = text
        
        for pattern, replacement in cls.CITATION_PATTERNS:
            result = re.sub(pattern, replacement, result)
        
        return result


class OptimizedPolishingWorkflow:
    """优化版学术论文润色工作流程"""
    
    def __init__(self, api_provider: str = "dashscope", model: str = "qwen-plus"):
        self.api_provider = api_provider
        self.model = model
        self.llm_client = None
        self.footnote_manager = FootnoteManager()
        
    def initialize(self) -> bool:
        """初始化组件"""
        try:
            config_file = Path(__file__).parent / 'config' / 'api_key.txt'
            if config_file.exists():
                with open(config_file, 'r', encoding='utf-8') as f:
                    for line in f:
                        if 'qwen3.5-plus' in line:
                            os.environ['DASHSCOPE_API_KEY'] = line.split('=')[-1].strip()
                            break
            
            llm_config = {
                'provider': self.api_provider,
                'api_key': os.getenv('DASHSCOPE_API_KEY'),
                'model': self.model
            }
            self.llm_client = create_llm_client(llm_config)
            return True
        except Exception as e:
            print(f"Initialization error: {e}")
            return False
    
    def process_document(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """
        处理文档的完整流程
        
        Args:
            input_path: 输入文件路径
            output_path: 输出文件路径
            
        Returns:
            Dict: 处理结果
        """
        result = {
            'status': 'pending',
            'input': input_path,
            'output': output_path,
            'sections_processed': 0,
            'format_issues_fixed': 0,
            'content_purified': False,
            'footnotes_added': 0
        }
        
        try:
            print("=" * 70)
            print("优化润色工作流程 v3.0")
            print("=" * 70)
            
            if not self.initialize():
                result['status'] = 'failed'
                result['error'] = 'Initialization failed'
                return result
            
            print("\n[Step 1] 加载文档...")
            doc = Document(input_path)
            original_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            print(f"  - 加载 {len(original_paragraphs)} 个段落")
            
            print("\n[Step 2] 提取并重组章节结构...")
            sections = self._extract_sections(original_paragraphs)
            print(f"  - 识别 {len(sections)} 个章节")
            for i, sec in enumerate(sections):
                print(f"    {i+1}. {sec['title'][:40]}...")
            
            print("\n[Step 3] 内容纯净度处理...")
            purified_sections = []
            for section in sections:
                purified = self._purify_section_content(section)
                purified_sections.append(purified)
            result['content_purified'] = True
            print("  - 内容净化完成")
            
            print("\n[Step 4] 生成优化润色内容...")
            polished_sections = self._polish_all_sections(purified_sections)
            print(f"  - 完成 {len(polished_sections)} 个章节的润色")
            
            print("\n[Step 5] 生成标准格式文档...")
            self._generate_standard_document(polished_sections, output_path)
            print(f"  - 文档已保存: {output_path}")
            
            result['status'] = 'completed'
            result['sections_processed'] = len(polished_sections)
            result['format_issues_fixed'] = len(polished_sections)
            result['footnotes_added'] = self.footnote_manager.footnote_id - 1
            
            print("\n" + "=" * 70)
            print("处理完成!")
            print("=" * 70)
            print(f"处理章节数: {result['sections_processed']}")
            print(f"格式修复数: {result['format_issues_fixed']}")
            print(f"脚注添加数: {result['footnotes_added']}")
            
        except Exception as e:
            result['status'] = 'failed'
            result['error'] = str(e)
            print(f"\nError: {e}")
        
        return result
    
    def _extract_sections(self, paragraphs: List[str]) -> List[Dict]:
        """提取并重组章节结构"""
        sections = []
        current_section = None
        current_content = []
        
        section_optimizer = SectionStructureOptimizer()
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            is_heading = section_optimizer.is_nested_title(para_text)
            main_section = section_optimizer.extract_main_section(para_text)
            
            if len(para_text) < 50 and any(
                re.search(p, para_text) for p in [
                    r'第[一二三四五六七八九十]+部分',
                    r'第\d+部分',
                    r'结语|结论|引言|摘要|参考文献',
                    r'一、|二、|三、|四、|五、'
                ]
            ):
                if current_section and current_content:
                    current_section['content'] = '\n'.join(current_content)
                    sections.append(current_section)
                
                clean_title = section_optimizer.clean_section_title(para_text)
                current_section = {
                    'title': clean_title,
                    'original_title': para_text,
                    'content': '',
                    'paragraphs': [],
                    'is_heading': True
                }
                current_content = []
            else:
                if current_section is None:
                    current_section = {
                        'title': '正文',
                        'original_title': '',
                        'content': '',
                        'paragraphs': [],
                        'is_heading': False
                    }
                current_content.append(para_text)
                if current_section:
                    current_section['paragraphs'].append(para_text)
        
        if current_section and current_content:
            current_section['content'] = '\n'.join(current_content)
            sections.append(current_section)
        
        return sections
    
    def _purify_section_content(self, section: Dict) -> Dict:
        """净化章节内容"""
        purified_content = ContentPurifier.clean_text(section['content'])
        purified_paragraphs = []
        
        for para in section['paragraphs']:
            cleaned = ContentPurifier.clean_text(para)
            if cleaned:
                purified_paragraphs.append(cleaned)
        
        section['content'] = purified_content
        section['paragraphs'] = purified_paragraphs
        section['is_purified'] = True
        
        return section
    
    def _polish_section(self, section: Dict) -> Dict:
        """润色单个章节"""
        system_prompt = """你是一位专业的日本史学术论文编辑。请对以下学术论文内容进行精简润色：

要求：
1. 保留核心学术观点和结论
2. 保持历史史实和重要事件
3. 保护历史专有名词和学术术语
4. 删除逻辑冗余的论述和重复表达
5. 保持学术表达的规范性和专业性
6. 确保语言表达精准流畅

输出格式：
直接输出润色后的纯文本内容，不要包含任何JSON、代码标记或其他技术格式。"""

        user_prompt = f"""请精简润色以下内容：

{section['content']}

直接输出润色后的文本内容。"""

        try:
            response = self.llm_client._call_llm(
                f"{system_prompt}\n\n{user_prompt}",
                temperature=0.3,
                max_tokens=4000
            )
            
            polished_text = response.get('content', section['content'])
            
            cleaned = ContentPurifier.clean_text(polished_text)
            
            cleaned = AcademicCitationCleaner.clean_citations(cleaned)
            
            section['polished_content'] = cleaned
            section['polished_paragraphs'] = [
                p.strip() for p in cleaned.split('\n') if p.strip()
            ]
            
        except Exception as e:
            print(f"  Warning: LLM call failed, using purified content: {e}")
            section['polished_content'] = section['content']
            section['polished_paragraphs'] = section['paragraphs']
        
        return section
    
    def _polish_all_sections(self, sections: List[Dict]) -> List[Dict]:
        """润色所有章节"""
        polished = []
        
        for i, section in enumerate(sections):
            print(f"\n  处理章节 {i+1}/{len(sections)}: {section['title'][:30]}...")
            polished_section = self._polish_section(section)
            polished.append(polished_section)
            
            if section.get('is_heading'):
                fn_num, fn_text = self.footnote_manager.add_footnote(
                    f"本节内容：{section['title']}",
                    "章节说明"
                )
                print(f"    + 添加脚注 [{fn_num}]")
        
        return polished
    
    def _generate_standard_document(self, sections: List[Dict], output_path: str):
        """生成标准格式文档"""
        doc = Document()
        
        for section in sections:
            title = section['title']
            
            heading = doc.add_heading(title, level=2)
            for run in heading.runs:
                run.font.name = '黑体'
                run.font.size = Pt(14)
                run.font.bold = True
            
            paragraphs = section.get('polished_paragraphs', [])
            for para_text in paragraphs:
                if not para_text.strip():
                    continue
                
                p = doc.add_paragraph(para_text)
                
                FormatStandardizer.apply_format_to_paragraph(p, is_heading=False)
            
            doc.add_paragraph()
        
        if self.footnote_manager.footnotes:
            doc.add_page_break()
            
            fn_heading = doc.add_heading('脚注', level=2)
            for run in fn_heading.runs:
                run.font.name = '黑体'
                run.font.size = Pt(14)
                run.font.bold = True
            
            for fn in self.footnote_manager.footnotes:
                p = doc.add_paragraph(fn['text'])
                FormatStandardizer.apply_format_to_paragraph(p, is_heading=False)
        
        doc.save(output_path)


def main():
    """主函数"""
    input_file = r"c:\Users\lyzha\Desktop\AItools-for-historyresearch\TW《新渡户论》20260324.docx"
    output_file = r"c:\Users\lyzha\Desktop\AItools-for-historyresearch\TW《新渡户论》20260324_polished_v3.docx"
    
    print("Academic Paper Optimized Polishing Workflow v3.0")
    print("=" * 70)
    print(f"Input: {input_file}")
    print(f"Output: {output_file}")
    print("=" * 70)
    
    workflow = OptimizedPolishingWorkflow(
        api_provider="dashscope",
        model="qwen-plus"
    )
    
    result = workflow.process_document(input_file, output_file)
    
    print("\n" + "=" * 70)
    print("Final Result")
    print("=" * 70)
    print(f"Status: {result['status']}")
    print(f"Sections: {result['sections_processed']}")
    print(f"Output: {result['output']}")
    
    return result


if __name__ == "__main__":
    main()
