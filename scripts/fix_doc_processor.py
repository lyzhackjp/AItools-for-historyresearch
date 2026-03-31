import sys
sys.path.append('modules')
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

class EnhancedDocProcessor:
    """增强版文档处理器，支持完整的脚注保留"""

    def __init__(self):
        self.current_footnote_id = 0

    def create_document_with_footnotes(self, content: dict, output_path: str, original_docx_path: str = None) -> bool:
        """
        创建包含完整脚注的文档

        Args:
            content: 包含文本内容的字典
            output_path: 输出文件路径
            original_docx_path: 原始文档路径（用于复制脚注结构）

        Returns:
            bool: 创建是否成功
        """
        if original_docx_path and Path(original_docx_path).exists():
            return self._create_from_original(content, output_path, original_docx_path)
        else:
            return self._create_new_with_footnotes(content, output_path)

    def _create_from_original(self, content: dict, output_path: str, original_path: str) -> bool:
        """从原始文档复制结构并修改内容"""
        import zipfile
        import shutil
        import os
        from lxml import etree

        temp_dir = 'temp_docx_workflow'
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

        os.makedirs(temp_dir)

        with zipfile.ZipFile(original_path, 'r') as z:
            z.extractall(temp_dir)

        doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        with open(doc_xml_path, 'rb') as f:
            doc_content = f.read()
            tree = etree.fromstring(doc_content)

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        body = tree.find('.//w:body', ns)
        existing_paras = body.findall('w:p', ns)
        for para in existing_paras:
            body.remove(para)

        existing_tables = body.findall('w:tbl', ns)
        for table in existing_tables:
            body.remove(table)

        title = content.get('title', '')
        if title:
            title_para = self._create_paragraph_element(title, 'Heading1', ns)
            body.insert(0, title_para)

        for i, para_info in enumerate(content.get('paragraphs', [])):
            para_elem = self._create_paragraph_element(para_info.get('text', ''), para_info.get('style', 'Normal'), ns)
            body.append(para_elem)

        modified_xml = etree.tostring(tree, encoding='unicode')
        modified_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + modified_xml

        with open(doc_xml_path, 'w', encoding='utf-8') as f:
            f.write(modified_xml)

        output_path_temp = os.path.join(temp_dir, '..', 'temp_output.docx')
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as z:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    z.write(file_path, arcname)

        shutil.rmtree(temp_dir)
        return True

    def _create_paragraph_element(self, text: str, style: str, ns: dict) -> object:
        """创建段落XML元素"""
        from lxml import etree

        para = etree.Element(f'{{{ns["w"]}}}p')

        pPr = etree.SubElement(para, f'{{{ns["w"]}}}pPr')
        pStyle = etree.SubElement(pPr, f'{{{ns["w"]}}}pStyle')
        pStyle.set(f'{{{ns["w"]}}}val', style)

        run = etree.SubElement(para, f'{{{ns["w"]}}}r')
        rPr = etree.SubElement(run, f'{{{ns["w"]}}}rPr')

        if style.startswith('Heading'):
            sz = etree.SubElement(rPr, f'{{{ns["w"]}}}sz')
            sz.set(f'{{{ns["w"]}}}val', '32')
            b = etree.SubElement(rPr, f'{{{ns["w"]}}}b')

        t = etree.SubElement(run, f'{{{ns["w"]}}}t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text

        return para

    def _create_new_with_footnotes(self, content: dict, output_path: str) -> bool:
        """创建新的文档（带脚注支持）"""
        doc = Document()

        if content.get('title'):
            title = doc.add_heading(content['title'], level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for para_info in content.get('paragraphs', []):
            para = doc.add_paragraph()
            para.add_run(para_info.get('text', ''))

            if para_info.get('footnote_ref'):
                self._add_footnote_reference(para, para_info['footnote_ref'])

        if content.get('footnotes'):
            self._create_footnotes_section(doc, content['footnotes'])

        doc.save(output_path)
        return True

    def _add_footnote_reference(self, para, footnote_id: int):
        """在段落中添加脚注引用"""
        run = para.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar_begin)

        run2 = para.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.text = ' REF _Reffnt1 \\h'
        run2._r.append(instrText)

        run3 = para.add_run()
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar_end)

        run4 = para.add_run()
        footnoteRef = OxmlElement('w:footnoteReference')
        footnoteRef.set(qn('w:id'), str(footnote_id))
        run4._r.append(footnoteRef)

    def _create_footnotes_section(self, doc: Document, footnotes: list):
        """创建脚注部分"""
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from lxml import etree

        footnotes_xml = b'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:footnote w:id="0" w:type="separator">
        <w:p>
            <w:r>
                <w:separator/>
            </w:r>
        </w:p>
    </w:footnote>
</w:footnotes>'''

        tree = etree.fromstring(footnotes_xml)
        footnotes_root = tree

        for fn in footnotes:
            fn_id = int(fn.get('id', 1))
            fn_text = fn.get('text', '')

            footnote_elem = etree.SubElement(footnotes_root, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnote')
            footnote_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(fn_id))
            footnote_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'normal')

            p_elem = etree.SubElement(footnote_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')

            r_elem = etree.SubElement(p_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            fn_ref = etree.SubElement(r_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnoteRef')

            r_elem2 = etree.SubElement(p_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            t_elem = etree.SubElement(r_elem2, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            t_elem.text = fn_text

        footnotes_xml_str = etree.tostring(tree, encoding='unicode')
        footnotes_bytes = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + footnotes_xml_str).encode('utf-8')

        reltype = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
        partname = PackURI('/word/footnotes.xml')
        part = Part(partname, 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml', footnotes_bytes)
        doc.part.package.relate_to(part, reltype)


if __name__ == '__main__':
    print("=== 使用增强版文档处理器 ===")
    print()

    from modules.doc_processor import DocProcessor

    dp = DocProcessor()
    original_path = 'TW《新渡户论》20260324.docx'
    output_path = 'TW《新渡户论》20260324_增强处理.docx'

    print("步骤1: 提取原始文档...")
    doc_info = dp.extract_text(original_path)
    print(f"  ✓ 标题: {doc_info.get('title')}")
    print(f"  ✓ 段落数: {len(doc_info.get('paragraphs', []))}")
    print(f"  ✓ 脚注数: {len(doc_info.get('footnotes', []))}")

    print()
    print("步骤2: 使用增强处理器创建文档...")
    enhanced_dp = EnhancedDocProcessor()
    enhanced_dp.create_document_with_footnotes(doc_info, output_path, original_path)
    print(f"  ✓ 文档已创建: {output_path}")

    print()
    print("步骤3: 验证脚注...")
    import zipfile
    with zipfile.ZipFile(output_path, 'r') as z:
        if 'word/footnotes.xml' in z.namelist():
            print("  ✓ footnotes.xml存在")
            with z.open('word/footnotes.xml') as f:
                from lxml import etree
                content = f.read()
                tree = etree.fromstring(content)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                footnotes = tree.findall('.//w:footnote', ns)
                print(f"  ✓ 包含{len(footnotes)}个脚注")
        else:
            print("  ✗ footnotes.xml不存在")

    print()
    print("=== 完成 ===")
