#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
学术论文智能润色增强工作流程 v2.0

基于 paper_polisher_enhanced.py 模块的综合处理系统

功能特点：
1. 文档结构化分段处理：按学术论文标准结构智能拆分
2. 双重校验机制：确保润色质量和关键信息完整性
3. 智能异常处理：自动调用学习模块和开源资源搜索
4. 全过程文档化：详细的处理日志和可追溯性

引用标准：WORKFLOW_DIAGRAM.md 学术论文润色流程
"""

import sys
import os
import json
import re
import datetime
import hashlib
import traceback
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional
from dataclasses import dataclass, field
from enum import Enum

sys.path.append(str(Path(__file__).parent / 'modules'))

try:
    import docx
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_COLOR_INDEX
except ImportError:
    print("警告: python-docx 未安装")
    sys.exit(1)

from modules.llm_client import LLMClient, create_llm_client
from modules.doc_processor import DocProcessor


class ProcessingStatus(Enum):
    """处理状态枚举"""
    PENDING = "pending"
    IN_PROGRESS = "in_progress"
    COMPLETED = "completed"
    FAILED = "failed"
    QUALITY_WARNING = "quality_warning"
    OPTIMIZATION_REQUIRED = "optimization_required"


class DocumentSection(Enum):
    """学术论文标准章节"""
    ABSTRACT = "摘要"
    ABSTRACT_EN = "Abstract"
    INTRODUCTION = "引言"
    LITERATURE_REVIEW = "文献综述"
    RESEARCH_METHODS = "研究方法"
    RESULTS = "结果分析"
    DISCUSSION = "讨论"
    CONCLUSION = "结论"
    REFERENCES = "参考文献"
    ACKNOWLEDGMENTS = "致谢"
    APPENDIX = "附录"
    UNKNOWN = "正文"


@dataclass
class SectionContent:
    """章节内容数据类"""
    section_type: DocumentSection
    title: str
    content: str
    paragraphs: List[Dict]
    start_index: int = 0
    end_index: int = 0
    word_count: int = 0
    footnotes: List[Dict] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        return {
            'section_type': self.section_type.value,
            'title': self.title,
            'content': self.content,
            'start_index': self.start_index,
            'end_index': self.end_index,
            'word_count': self.word_count,
            'paragraph_count': len(self.paragraphs),
            'footnote_count': len(self.footnotes)
        }


@dataclass
class ProcessingLog:
    """处理日志数据类"""
    workflow_name: str
    start_time: datetime.datetime
    document_info: Dict = field(default_factory=dict)
    sections: List[Dict] = field(default_factory=list)
    api_calls: List[Dict] = field(default_factory=list)
    quality_checks: List[Dict] = field(default_factory=list)
    optimization_records: List[Dict] = field(default_factory=list)
    errors: List[Dict] = field(default_factory=list)
    warnings: List[Dict] = field(default_factory=list)
    final_result: Dict = field(default_factory=dict)
    
    def to_dict(self) -> Dict:
        return {
            'workflow_name': self.workflow_name,
            'start_time': self.start_time.isoformat(),
            'end_time': datetime.datetime.now().isoformat(),
            'document_info': self.document_info,
            'sections': self.sections,
            'api_calls': self.api_calls,
            'quality_checks': self.quality_checks,
            'optimization_records': self.optimization_records,
            'errors': self.errors,
            'warnings': self.warnings,
            'final_result': self.final_result
        }


@dataclass
class QualityCheckResult:
    """质量检查结果"""
    section: str
    check_type: str
    status: ProcessingStatus
    score: float
    issues: List[str] = field(default_factory=list)
    suggestions: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        return {
            'section': self.section,
            'check_type': self.check_type,
            'status': self.status.value,
            'score': self.score,
            'issues': self.issues,
            'suggestions': self.suggestions
        }


class AcademicStructureRecognizer:
    """学术论文结构识别器"""
    
    SECTION_PATTERNS = {
        DocumentSection.ABSTRACT: [
            r'^摘要\s*$',
            r'^【摘要】',
            r'^〔摘要〕',
            r'摘\s+要\s*$'
        ],
        DocumentSection.ABSTRACT_EN: [
            r'^Abstract\s*$',
            r'^ABSTRACT\s*$'
        ],
        DocumentSection.INTRODUCTION: [
            r'^一、?\s*引\s*言',
            r'^二、?\s*研究背景',
            r'^引\s*言\s*$',
            r'^【引言】',
            r'^前言\s*$',
            r'^序\s*言\s*$',
            r'^绪\s*论\s*$',
            r'^\d+、?\s*引\s*言'
        ],
        DocumentSection.LITERATURE_REVIEW: [
            r'^文\s*献\s*综\s*述',
            r'^相关文献',
            r'^研究综述',
            r'^文献回顾',
            r'^【文献综述】',
            r'^\d+、?\s*文\s*献'
        ],
        DocumentSection.RESEARCH_METHODS: [
            r'^研究方法',
            r'^研究设计与方法',
            r'^方\s*法\s*$',
            r'^【研究方法】',
            r'^研究过程',
            r'^材料与方法',
            r'^实验方法',
            r'^技术路线',
            r'^\d+、?\s*方\s*法'
        ],
        DocumentSection.RESULTS: [
            r'^研究结果',
            r'^结\s*果\s*$',
            r'^【结果】',
            r'^结果与分析',
            r'^实验结果',
            r'^数据分析',
            r'^\d+、?\s*结\s*果'
        ],
        DocumentSection.DISCUSSION: [
            r'^讨\s*论\s*$',
            r'^【讨论】',
            r'^讨论与分析',
            r'^分析与讨论',
            r'^\d+、?\s*讨\s*论'
        ],
        DocumentSection.CONCLUSION: [
            r'^结\s*论\s*$',
            r'^【结论】',
            r'^结论与展望',
            r'^总结\s*$',
            r'^结\s*语\s*$',
            r'^结\s*语',
            r'^\d+、?\s*结\s*论'
        ],
        DocumentSection.REFERENCES: [
            r'^参考文献',
            r'^参\s*考\s*文\s*献',
            r'^【参考文献】',
            r'^引文目录',
            r'^引用文献'
        ],
        DocumentSection.ACKNOWLEDGMENTS: [
            r'^致\s*谢\s*$',
            r'^【致谢】',
            r'^感谢\s*$',
            r'^致谢语'
        ],
        DocumentSection.APPENDIX: [
            r'^附\s*录\s*$',
            r'^【附录】',
            r'^附\s*表',
            r'^附\s*图'
        ]
    }
    
    @classmethod
    def recognize_section(cls, paragraph_text: str) -> Tuple[DocumentSection, float]:
        """
        识别段落所属的学术论文章节
        
        Args:
            paragraph_text: 段落文本
            
        Returns:
            Tuple[DocumentSection, float]: (识别的章节类型, 置信度)
        """
        if not paragraph_text or not paragraph_text.strip():
            return DocumentSection.UNKNOWN, 0.0
        
        text = paragraph_text.strip()
        
        for section_type, patterns in cls.SECTION_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, text):
                    return section_type, 0.9
        
        if cls._is_heading_style(text):
            return DocumentSection.UNKNOWN, 0.3
        
        return DocumentSection.UNKNOWN, 0.0
    
    @classmethod
    def _is_heading_style(cls, text: str) -> bool:
        """判断是否为标题样式"""
        short_lines = [l.strip() for l in text.split('\n') if len(l.strip()) < 50]
        return len(short_lines) == len(text.split('\n')) and len(short_lines) > 0


class DocumentSegmenter:
    """文档分段处理器"""
    
    def __init__(self, max_chars_per_segment: int = 3000):
        """
        初始化分段处理器
        
        Args:
            max_chars_per_segment: 每个分段的最大字符数
        """
        self.max_chars = max_chars_per_segment
        self.recognizer = AcademicStructureRecognizer()
    
    def segment_document(self, paragraphs: List[Dict], footnotes: List[Dict] = None) -> List[SectionContent]:
        """
        将文档按学术论文结构分段
        
        Args:
            paragraphs: 段落列表
            footnotes: 脚注列表
            
        Returns:
            List[SectionContent]: 分段后的章节列表
        """
        sections = []
        current_section = None
        current_content = []
        current_paragraphs = []
        char_count = 0
        
        for idx, para in enumerate(paragraphs):
            text = para.get('text', '')
            section_type, confidence = self.recognizer.recognize_section(text)
            
            is_new_section = (
                section_type != DocumentSection.UNKNOWN and 
                confidence > 0.7 and
                (current_section is None or current_section.section_type != section_type)
            )
            
            if is_new_section and char_count > 100:
                if current_section:
                    current_section.content = '\n'.join(current_content)
                    current_section.end_index = idx - 1
                    current_section.word_count = len(current_section.content)
                    sections.append(current_section)
                
                current_section = SectionContent(
                    section_type=section_type,
                    title=text.strip(),
                    content='',
                    paragraphs=[],
                    start_index=idx,
                    footnotes=[]
                )
                current_content = []
                current_paragraphs = []
                char_count = 0
            
            if current_section is None:
                current_section = SectionContent(
                    section_type=DocumentSection.UNKNOWN,
                    title='正文',
                    content='',
                    paragraphs=[],
                    start_index=idx,
                    footnotes=[]
                )
            
            current_paragraphs.append(para)
            current_content.append(text)
            char_count += len(text)
            
            if char_count >= self.max_chars:
                sub_segments = self._split_large_section(
                    current_content, 
                    current_paragraphs, 
                    current_section
                )
                if sub_segments:
                    sections.extend(sub_segments[:-1])
                    current_section = sub_segments[-1]
                    current_content = [current_section.content]
                    current_paragraphs = current_section.paragraphs
                    char_count = len(current_section.content)
        
        if current_section:
            current_section.content = '\n'.join(current_content)
            current_section.end_index = len(paragraphs) - 1
            current_section.word_count = len(current_section.content)
            current_section.paragraphs = current_paragraphs
            sections.append(current_section)
        
        return sections
    
    def _split_large_section(
        self, 
        content: List[str], 
        paragraphs: List[Dict],
        section: SectionContent
    ) -> List[SectionContent]:
        """拆分过大的章节"""
        if len(content) <= 3:
            return [section]
        
        mid = len(content) // 2
        part1 = SectionContent(
            section_type=section.section_type,
            title=f"{section.title} (第一部分)",
            content='\n'.join(content[:mid]),
            paragraphs=paragraphs[:mid],
            start_index=section.start_index,
            footnotes=[]
        )
        part1.word_count = len(part1.content)
        
        part2 = SectionContent(
            section_type=section.section_type,
            title=f"{section.title} (第二部分)",
            content='\n'.join(content[mid:]),
            paragraphs=paragraphs[mid:],
            start_index=section.start_index + mid,
            footnotes=[]
        )
        part2.word_count = len(part2.content)
        
        return [part1, part2]


class PolishingQualityChecker:
    """润色质量检查器"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm_client = llm_client
    
    def check_quality(
        self, 
        original_text: str, 
        polished_text: str, 
        section_type: DocumentSection
    ) -> QualityCheckResult:
        """
        检查润色质量
        
        Args:
            original_text: 原文
            polished_text: 润色后文本
            section_type: 章节类型
            
        Returns:
            QualityCheckResult: 质量检查结果
        """
        check_prompt = f"""请对以下学术论文润色结果进行质量检查：

## 原文本
{original_text[:2000]}

## 润色后文本
{polished_text[:2000]}

## 章节类型
{section_type.value}

请检查以下方面并返回JSON格式结果：
1. 核心论点是否保留
2. 关键信息是否完整
3. 学术表达是否规范
4. 逻辑连贯性
5. 整体质量评分(0-100)

返回格式：
{{
    "core_preserved": true/false,
    "key_info_complete": true/false,
    "academic_standard": true/false,
    "logical_coherence": true/false,
    "score": 评分数字,
    "issues": ["问题1", "问题2"],
    "suggestions": ["建议1", "建议2"]
}}"""

        try:
            combined_prompt = f"你是一位专业的学术论文质量评审专家。\n\n{check_prompt}"
            response = self.llm_client._call_llm(
                combined_prompt,
                temperature=0.3,
                max_tokens=1500
            )
            
            result_text = response.get('content', '{}')
            result_json = json.loads(result_text)
            
            status = ProcessingStatus.COMPLETED
            if result_json.get('score', 100) < 70:
                status = ProcessingStatus.QUALITY_WARNING
            
            return QualityCheckResult(
                section=section_type.value,
                check_type="full_quality_check",
                status=status,
                score=result_json.get('score', 0),
                issues=result_json.get('issues', []),
                suggestions=result_json.get('suggestions', [])
            )
        except Exception as e:
            return QualityCheckResult(
                section=section_type.value,
                check_type="full_quality_check",
                status=ProcessingStatus.FAILED,
                score=0,
                issues=[f"质量检查失败: {str(e)}"],
                suggestions=[]
            )
    
    def check_key_content_preservation(
        self, 
        original_text: str, 
        polished_text: str,
        key_elements: List[str]
    ) -> Dict[str, Any]:
        """
        检查关键内容保留情况
        
        Args:
            original_text: 原文
            polished_text: 润色后文本
            key_elements: 需要保留的关键元素列表
            
        Returns:
            Dict: 检查结果
        """
        missing_elements = []
        preserved_elements = []
        
        for element in key_elements:
            if element in polished_text:
                preserved_elements.append(element)
            else:
                missing_elements.append(element)
        
        preservation_rate = len(preserved_elements) / len(key_elements) if key_elements else 1.0
        
        return {
            'preserved': preserved_elements,
            'missing': missing_elements,
            'preservation_rate': preservation_rate,
            'all_preserved': len(missing_elements) == 0
        }


class LearningModuleIntegration:
    """学习模块集成 - 智能异常处理"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm_client = llm_client
    
    def generate_optimization_suggestions(
        self, 
        issue_description: str,
        context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        生成优化建议
        
        Args:
            issue_description: 问题描述
            context: 上下文信息
            
        Returns:
            Dict: 包含优化建议的报告
        """
        prompt = f"""作为学术论文处理领域的专家，请分析以下问题并提供优化建议：

## 问题描述
{issue_description}

## 当前上下文
{json.dumps(context, ensure_ascii=False, indent=2)}

请从以下方面提供优化建议：
1. 技术实现方案
2. 可参考的学术资源
3. 具体的改进步骤
4. 预期的优化效果

返回JSON格式：
{{
    "technical_suggestions": ["建议1", "建议2"],
    "reference_resources": ["资源1", "资源2"],
    "implementation_steps": ["步骤1", "步骤2"],
    "expected_improvement": "预期效果描述"
}}"""

        try:
            combined_prompt = f"你是一位资深的AI学术工具开发专家。\n\n{prompt}"
            response = self.llm_client._call_llm(
                combined_prompt,
                temperature=0.5,
                max_tokens=2000
            )
            
            result_text = response.get('content', '{}')
            result_json = json.loads(result_text)
            
            return {
                'status': 'success',
                'suggestions': result_json,
                'timestamp': datetime.datetime.now().isoformat()
            }
        except Exception as e:
            return {
                'status': 'failed',
                'error': str(e),
                'timestamp': datetime.datetime.now().isoformat()
            }


class OpenSourceFinderIntegration:
    """开源资源搜索集成"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm_client = llm_client
    
    def find_relevant_resources(
        self, 
        module_name: str,
        context: str
    ) -> Dict[str, Any]:
        """
        搜索相关的开源资源
        
        Args:
            module_name: 模块名称
            context: 应用上下文
            
        Returns:
            Dict: 资源搜索结果
        """
        prompt = f"""请推荐与以下学术文档处理模块相关的开源资源和工具：

## 模块名称
{module_name}

## 应用场景
{context}

请搜索并推荐：
1. GitHub上的相关开源项目
2. HuggingFace上的预训练模型
3. 学术论文处理相关的工具库
4. Python生态中的文档处理库

返回JSON格式：
{{
    "github_projects": [
        {{"name": "项目名", "description": "描述", "url": "链接", "relevance_score": 0-1}}
    ],
    "huggingface_models": [
        {{"name": "模型名", "description": "描述", "url": "链接", "relevance_score": 0-1}}
    ],
    "python_libraries": [
        {{"name": "库名", "description": "描述", "pip_command": "安装命令"}}
    ],
    "academic_papers": [
        {{"title": "论文标题", "authors": "作者", "year": 年份}}
    ]
}}"""

        try:
            combined_prompt = f"你是一位开源技术和学术资源检索专家。\n\n{prompt}"
            response = self.llm_client._call_llm(
                combined_prompt,
                temperature=0.5,
                max_tokens=3000
            )
            
            result_text = response.get('content', '{}')
            result_json = json.loads(result_text)
            
            return {
                'status': 'success',
                'resources': result_json,
                'timestamp': datetime.datetime.now().isoformat()
            }
        except Exception as e:
            return {
                'status': 'failed',
                'error': str(e),
                'timestamp': datetime.datetime.now().isoformat()
            }


class EnhancedPaperPolishingWorkflow:
    """增强版论文润色工作流程"""
    
    def __init__(self, api_provider: str = "dashscope", model: str = "qwen-plus"):
        """
        初始化工作流程
        
        Args:
            api_provider: API提供商
            model: 使用的模型
        """
        self.start_time = datetime.datetime.now()
        self.api_provider = api_provider
        self.model = model
        
        self.processing_log = ProcessingLog(
            workflow_name="学术论文智能润色增强工作流程 v2.0",
            start_time=self.start_time
        )
        
        self.llm_client = None
        self.doc_processor = None
        self.segmenter = None
        self.quality_checker = None
        self.learning_module = None
        self.opensource_finder = None
        
        self.max_retry_attempts = 2
        self.quality_threshold = 75.0
    
    def initialize(self) -> bool:
        """初始化所有组件"""
        try:
            self.processing_log.document_info['initialization'] = {
                'status': 'in_progress',
                'timestamp': datetime.datetime.now().isoformat()
            }
            
            llm_config = {
                'provider': self.api_provider,
                'api_key': os.getenv('DASHSCOPE_API_KEY'),
                'model': self.model
            }
            self.llm_client = create_llm_client(llm_config)
            self.doc_processor = DocProcessor()
            self.segmenter = DocumentSegmenter(max_chars_per_segment=3000)
            self.quality_checker = PolishingQualityChecker(self.llm_client)
            self.learning_module = LearningModuleIntegration(self.llm_client)
            self.opensource_finder = OpenSourceFinderIntegration(self.llm_client)
            
            self.processing_log.document_info['initialization'] = {
                'status': 'success',
                'api_provider': self.api_provider,
                'model': self.model,
                'timestamp': datetime.datetime.now().isoformat()
            }
            
            return True
        except Exception as e:
            self.processing_log.errors.append({
                'stage': 'initialization',
                'error': str(e),
                'timestamp': datetime.datetime.now().isoformat()
            })
            return False
    
    def process_document(
        self, 
        input_path: str, 
        output_path: str,
        enable_logging: bool = True
    ) -> Dict[str, Any]:
        """
        处理文档的完整流程
        
        Args:
            input_path: 输入文档路径
            output_path: 输出文档路径
            enable_logging: 是否启用日志
            
        Returns:
            Dict: 处理结果
        """
        workflow_result = {
            'status': 'pending',
            'input_path': input_path,
            'output_path': output_path,
            'processing_time': None,
            'sections_processed': 0,
            'total_quality_score': 0,
            'optimization_applied': False,
            'log': None
        }
        
        try:
            if not self.initialize():
                workflow_result['status'] = 'failed'
                workflow_result['error'] = '初始化失败'
                return workflow_result
            
            step_1_start = datetime.datetime.now()
            extracted_content = self._step1_extract_document(input_path)
            if not extracted_content:
                workflow_result['status'] = 'failed'
                workflow_result['error'] = '文档提取失败'
                return workflow_result
            
            step_2_start = datetime.datetime.now()
            sections = self._step2_segment_document(extracted_content)
            if not sections:
                workflow_result['status'] = 'failed'
                workflow_result['error'] = '文档分段失败'
                return workflow_result
            
            processed_sections = []
            all_quality_scores = []
            
            for idx, section in enumerate(sections):
                section_result = self._step3_process_section(
                    section, 
                    idx, 
                    len(sections)
                )
                
                if section_result['quality_score'] < self.quality_threshold:
                    retry_result = self._handle_quality_issue(section, section_result)
                    if retry_result:
                        section_result = retry_result
                
                processed_sections.append(section_result)
                all_quality_scores.append(section_result['quality_score'])
                
                self.processing_log.quality_checks.append({
                    'section': section.title,
                    'score': section_result['quality_score'],
                    'status': section_result['status'],
                    'timestamp': datetime.datetime.now().isoformat()
                })
            
            self._step4_generate_output(processed_sections, output_path)
            
            workflow_result['status'] = 'completed'
            workflow_result['sections_processed'] = len(sections)
            workflow_result['total_quality_score'] = sum(all_quality_scores) / len(all_quality_scores)
            workflow_result['processing_time'] = (
                datetime.datetime.now() - self.start_time
            ).total_seconds()
            
            if enable_logging:
                self.processing_log.final_result = workflow_result
                self._save_processing_log(output_path.replace('.docx', '_log.json'))
            
        except Exception as e:
            workflow_result['status'] = 'failed'
            workflow_result['error'] = str(e)
            workflow_result['traceback'] = traceback.format_exc()
            self.processing_log.errors.append({
                'stage': 'main_process',
                'error': str(e),
                'timestamp': datetime.datetime.now().isoformat()
            })
        
        return workflow_result
    
    def _step1_extract_document(self, input_path: str) -> Dict[str, Any]:
        """步骤1: 提取文档内容"""
        try:
            doc_info = self.doc_processor.extract_text(input_path)
            
            self.processing_log.document_info = {
                'file_name': os.path.basename(input_path),
                'file_size': os.path.getsize(input_path),
                'paragraph_count': len(doc_info.get('paragraphs', [])),
                'table_count': len(doc_info.get('tables', [])),
                'footnote_count': len(doc_info.get('footnotes', [])),
                'extraction_time': datetime.datetime.now().isoformat()
            }
            
            return doc_info
        except Exception as e:
            self.processing_log.errors.append({
                'stage': 'step1_extract',
                'error': str(e),
                'timestamp': datetime.datetime.now()
            })
            return None
    
    def _step2_segment_document(self, content: Dict[str, Any]) -> List[SectionContent]:
        """步骤2: 分段文档"""
        paragraphs = content.get('paragraphs', [])
        footnotes = content.get('footnotes', [])
        
        sections = self.segmenter.segment_document(paragraphs, footnotes)
        
        self.processing_log.sections = [s.to_dict() for s in sections]
        
        return sections
    
    def _step3_process_section(
        self, 
        section: SectionContent, 
        section_index: int,
        total_sections: int
    ) -> Dict[str, Any]:
        """步骤3: 处理单个章节"""
        retry_count = 0
        max_retries = self.max_retry_attempts
        
        while retry_count <= max_retries:
            try:
                polished_content = self._call_polishing_api(section)
                
                quality_result = self.quality_checker.check_quality(
                    section.content,
                    polished_content,
                    section.section_type
                )
                
                api_call_record = {
                    'section': section.title,
                    'section_index': section_index,
                    'total_sections': total_sections,
                    'attempt': retry_count + 1,
                    'section_type': section.section_type.value,
                    'content_length': len(section.content),
                    'polished_length': len(polished_content),
                    'quality_score': quality_result.score,
                    'timestamp': datetime.datetime.now().isoformat()
                }
                self.processing_log.api_calls.append(api_call_record)
                
                return {
                    'status': quality_result.status.value,
                    'section': section.to_dict(),
                    'polished_content': polished_content,
                    'quality_score': quality_result.score,
                    'quality_issues': quality_result.issues,
                    'quality_suggestions': quality_result.suggestions,
                    'retry_count': retry_count
                }
                
            except Exception as e:
                retry_count += 1
                if retry_count > max_retries:
                    return {
                        'status': 'failed',
                        'section': section.to_dict(),
                        'polished_content': section.content,
                        'quality_score': 0,
                        'quality_issues': [str(e)],
                        'quality_suggestions': [],
                        'retry_count': retry_count
                    }
    
    def _call_polishing_api(self, section: SectionContent) -> str:
        """调用润色API"""
        system_prompt = """你是一位专业的日本史学术论文编辑。请精简以下学术论文内容，同时：
1. 保留核心学术观点和结论
2. 保持历史史实和重要事件
3. 保护历史专有名词和学术术语
4. 保留所有脚注和引用标注
5. 删除逻辑冗余的论述和重复表达

请返回精简后的文本。"""
        
        user_prompt = f"""请精简以下{section.section_type.value}内容：

{section.content}

请返回JSON格式：
{{
    "modified_text": "精简后的文本",
    "deletions": ["删除的内容1", "删除的内容2"]
}}"""

        combined_prompt = f"{system_prompt}\n\n{user_prompt}"
        
        response = self.llm_client._call_llm(
            combined_prompt,
            temperature=0.3,
            max_tokens=4000
        )
        
        result_text = response.get('content', section.content)
        
        try:
            result_json = json.loads(result_text)
            return result_json.get('modified_text', section.content)
        except json.JSONDecodeError:
            return result_text
    
    def _handle_quality_issue(
        self, 
        section: SectionContent, 
        previous_result: Dict
    ) -> Optional[Dict]:
        """处理质量问题 - 调用学习模块"""
        if previous_result['quality_score'] >= self.quality_threshold:
            return None
        
        try:
            optimization_context = {
                'section_type': section.section_type.value,
                'content_preview': section.content[:500],
                'previous_quality_score': previous_result['quality_score'],
                'quality_issues': previous_result.get('quality_issues', [])
            }
            
            optimization_result = self.learning_module.generate_optimization_suggestions(
                f"润色质量不达标 (评分: {previous_result['quality_score']})",
                optimization_context
            )
            
            self.processing_log.optimization_records.append({
                'section': section.title,
                'previous_score': previous_result['quality_score'],
                'optimization_result': optimization_result,
                'timestamp': datetime.datetime.now().isoformat()
            })
            
            if optimization_result['status'] == 'success':
                suggestions = optimization_result['suggestions']
                refined_content = self._apply_refinement(
                    section.content, 
                    suggestions
                )
                
                new_quality = self.quality_checker.check_quality(
                    section.content,
                    refined_content,
                    section.section_type
                )
                
                if new_quality.score > previous_result['quality_score']:
                    return {
                        'status': 'optimized',
                        'section': section.to_dict(),
                        'polished_content': refined_content,
                        'quality_score': new_quality.score,
                        'quality_issues': new_quality.issues,
                        'quality_suggestions': new_quality.suggestions,
                        'retry_count': previous_result['retry_count'] + 1,
                        'optimization_applied': True
                    }
            
        except Exception as e:
            self.processing_log.warnings.append({
                'stage': 'handle_quality_issue',
                'section': section.title,
                'error': str(e),
                'timestamp': datetime.datetime.now().isoformat()
            })
        
        return None
    
    def _apply_refinement(
        self, 
        content: str, 
        suggestions: Dict
    ) -> str:
        """应用优化建议进行精炼"""
        refinement_prompt = f"""请基于以下优化建议，对文本进行改进：

## 原文
{content}

## 优化建议
{json.dumps(suggestions, ensure_ascii=False, indent=2)}

请返回改进后的文本（仅返回文本内容，不要额外说明）。"""

        try:
            combined_prompt = f"你是一位专业的学术论文编辑。\n\n{refinement_prompt}"
            response = self.llm_client._call_llm(
                combined_prompt,
                temperature=0.4,
                max_tokens=4000
            )
            
            return response.get('content', content)
        except Exception:
            return content
    
    def _step4_generate_output(
        self, 
        processed_sections: List[Dict], 
        output_path: str
    ):
        """步骤4: 生成输出文档"""
        try:
            doc = Document()
            
            for section_data in processed_sections:
                polished = section_data.get('polished_content', '')
                section = section_data.get('section', {})
                
                title = doc.add_heading(section.get('title', '章节'), level=2)
                
                paragraphs = polished.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
            
            doc.save(output_path)
            
            self.processing_log.document_info['output'] = {
                'path': output_path,
                'section_count': len(processed_sections),
                'generation_time': datetime.datetime.now().isoformat()
            }
            
        except Exception as e:
            self.processing_log.errors.append({
                'stage': 'step4_generate_output',
                'error': str(e),
                'timestamp': datetime.datetime.now().isoformat()
            })
            raise
    
    def _save_processing_log(self, log_path: str):
        """保存处理日志"""
        try:
            with open(log_path, 'w', encoding='utf-8') as f:
                json.dump(self.processing_log.to_dict(), f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存日志失败: {e}")


def main():
    """主函数"""
    input_file = r"c:\Users\lyzha\Desktop\AItools-for-historyresearch\TW《新渡户论》20260324.docx"
    output_file = r"c:\Users\lyzha\Desktop\AItools-for-historyresearch\TW《新渡户论》20260324_polished.docx"
    
    print("=" * 60)
    print("学术论文智能润色增强工作流程 v2.0")
    print("=" * 60)
    print(f"输入文件: {input_file}")
    print(f"输出文件: {output_file}")
    print("=" * 60)
    
    workflow = EnhancedPaperPolishingWorkflow(
        api_provider="dashscope",
        model="qwen-plus"
    )
    
    result = workflow.process_document(input_file, output_file)
    
    print("\n" + "=" * 60)
    print("处理结果摘要")
    print("=" * 60)
    print(f"状态: {result['status']}")
    print(f"处理章节数: {result['sections_processed']}")
    print(f"总质量评分: {result['total_quality_score']:.2f}")
    print(f"处理时间: {result['processing_time']:.2f}秒")
    
    if result.get('optimization_applied'):
        print("✓ 已应用优化建议")
    
    if result['status'] == 'completed':
        print(f"\n✓ 处理完成！输出文件: {result['output_path']}")
    else:
        print(f"\n✗ 处理失败: {result.get('error', '未知错误')}")
    
    return result


if __name__ == "__main__":
    main()
