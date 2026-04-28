"""
论文润色模块单元测试和集成测试

测试内容：
1. LLM客户端初始化测试
2. 段落润色功能测试
3. 修订模式应用测试
4. 脚注引用保护测试
5. 完整文档处理集成测试
"""

import os
import sys
import unittest
import tempfile
import shutil
import zipfile
from pathlib import Path

TEST_DIR = Path(__file__).parent
PROJECT_ROOT = TEST_DIR.parent
sys.path.insert(0, str(PROJECT_ROOT))

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ModuleNotFoundError as exc:
    raise unittest.SkipTest("python-docx is not installed") from exc

from config.api_key_manager import load_all_api_keys
load_all_api_keys()

from modules.paper_polisher import PaperPolisher, create_paper_polisher
from modules.llm_client import create_llm_client
from lxml import etree


class TestLLMClient(unittest.TestCase):
    """LLM客户端测试"""

    def test_create_llm_client_with_dict_config(self):
        """测试使用字典配置创建LLM客户端"""
        client = create_llm_client({'provider': 'dashscope'})
        self.assertIsNotNone(client)
        self.assertEqual(client.provider, 'dashscope')

    def test_llm_client_call(self):
        """测试LLM客户端调用"""
        client = create_llm_client({'provider': 'dashscope'})
        result = client._call_llm("请说'测试成功'", temperature=0.3, max_tokens=100)
        self.assertIn('content', result)
        self.assertIn('测试成功', result['content'])


class TestPaperPolisherInit(unittest.TestCase):
    """论文润色器初始化测试"""

    def test_create_paper_polisher_default(self):
        """测试默认参数创建润色器"""
        polisher = create_paper_polisher()
        self.assertIsNotNone(polisher)
        self.assertEqual(polisher.api_provider, 'qwen')

    def test_create_paper_polisher_with_provider(self):
        """测试指定provider创建润色器"""
        polisher = create_paper_polisher('minimax')
        self.assertEqual(polisher.api_provider, 'minimax')

    def test_min_paragraph_length(self):
        """测试最小段落长度设置"""
        polisher = create_paper_polisher()
        self.assertEqual(polisher.MIN_PARAGRAPH_LENGTH, 30)


class TestPolishParagraph(unittest.TestCase):
    """段落润色功能测试"""

    def setUp(self):
        self.polisher = create_paper_polisher()

    def test_polish_empty_paragraph(self):
        """测试空段落处理"""
        text, deletions = self.polisher.polish_paragraph("")
        self.assertEqual(text, "")
        self.assertEqual(deletions, [])

    def test_polish_short_paragraph(self):
        """测试短段落处理（低于最小长度）"""
        text, deletions = self.polisher.polish_paragraph("这是一个短段落。")
        self.assertEqual(text, "这是一个短段落。")
        self.assertEqual(deletions, [])

    def test_polish_normal_paragraph(self):
        """测试正常段落润色"""
        long_text = "近代日本思想中有一个非常有趣的现象，不少信奉来自西方的基督教、以普遍主义为志向的思想者，反而高度肯定作为日本特有之物的武士道。"
        text, deletions = self.polisher.polish_paragraph(long_text)
        self.assertIsNotNone(text)
        self.assertTrue(len(text) > 0)


class TestTrackChanges(unittest.TestCase):
    """修订模式功能测试"""

    def setUp(self):
        self.polisher = create_paper_polisher()
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.temp_dir)

    def test_enable_track_revisions(self):
        """测试启用修订追踪设置"""
        doc = Document()
        self.polisher._enable_track_revisions(doc)

        settings = doc.settings.element
        track_revisions = settings.find(qn('w:trackRevisions'))
        self.assertIsNotNone(track_revisions)

    def test_apply_track_changes_creates_elements(self):
        """测试修订标记元素创建"""
        doc = Document()
        para = doc.add_paragraph("原始文本")

        self.polisher._apply_track_changes(para, "原始文本", "修改后文本")

        del_elem = para._p.find(qn('w:del'))
        ins_elem = para._p.find(qn('w:ins'))

        self.assertIsNotNone(del_elem)
        self.assertIsNotNone(ins_elem)

    def test_apply_track_changes_preserves_footnote_refs(self):
        """测试修订标记保留脚注引用"""
        doc = Document()
        para = doc.add_paragraph()

        fn_run = OxmlElement('w:r')
        fn_ref = OxmlElement('w:footnoteReference')
        fn_ref.set(qn('w:id'), '1')
        fn_run.append(fn_ref)
        para._p.append(fn_run)

        self.polisher._apply_track_changes(para, "原始文本", "修改后文本")

        fn_refs = para._p.findall('.//w:footnoteReference',
                                  {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        self.assertEqual(len(fn_refs), 1)


class TestFootnoteProtection(unittest.TestCase):
    """脚注引用保护测试"""

    def setUp(self):
        self.polisher = create_paper_polisher()
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.temp_dir)

    def test_footnote_extraction(self):
        """测试脚注引用元素提取"""
        doc = Document()
        para = doc.add_paragraph()

        for i in range(1, 4):
            fn_run = OxmlElement('w:r')
            fn_ref = OxmlElement('w:footnoteReference')
            fn_ref.set(qn('w:id'), str(i))
            fn_run.append(fn_ref)
            para._p.append(fn_run)

        footnote_refs = []
        for child in para._p:
            if child.tag == qn('w:r'):
                fn_ref = child.find(qn('w:footnoteReference'))
                if fn_ref is not None:
                    footnote_refs.append(child)

        self.assertEqual(len(footnote_refs), 3)


class TestDocumentProcessing(unittest.TestCase):
    """完整文档处理集成测试"""

    def setUp(self):
        self.polisher = create_paper_polisher()
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.temp_dir)

    def test_process_simple_document(self):
        """测试处理简单文档"""
        input_path = Path(self.temp_dir) / 'input.docx'
        output_path = Path(self.temp_dir) / 'output.docx'

        doc = Document()
        doc.add_paragraph("近代日本思想中有一个非常有趣的现象，不少信奉来自西方的基督教、以普遍主义为志向的思想者，反而高度肯定作为日本特有之物的武士道。")
        doc.add_paragraph("这是另一个段落，用于测试文档处理功能。")
        doc.save(str(input_path))

        result = self.polisher.process_document(
            str(input_path),
            str(output_path),
            enable_track_changes=True
        )

        self.assertTrue(result['success'])
        self.assertTrue(output_path.exists())

    def test_process_document_with_footnotes(self):
        """测试处理包含脚注的文档"""
        input_path = Path(self.temp_dir) / 'input_fn.docx'
        output_path = Path(self.temp_dir) / 'output_fn.docx'

        doc = Document()
        para = doc.add_paragraph("测试文本")

        fn_run = OxmlElement('w:r')
        fn_ref = OxmlElement('w:footnoteReference')
        fn_ref.set(qn('w:id'), '1')
        fn_run.append(fn_ref)
        para._p.append(fn_run)

        doc.save(str(input_path))

        result = self.polisher.process_document(
            str(input_path),
            str(output_path),
            enable_track_changes=True
        )

        self.assertTrue(result['success'])

        with zipfile.ZipFile(str(output_path), 'r') as z:
            with z.open('word/document.xml') as f:
                tree = etree.parse(f)
                root = tree.getroot()
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                fn_refs = root.findall('.//w:footnoteReference', ns)
                self.assertEqual(len(fn_refs), 1)


class TestResponseCleaning(unittest.TestCase):
    """响应清理功能测试"""

    def setUp(self):
        self.polisher = create_paper_polisher()

    def test_clean_response_with_prefix(self):
        """测试清理带前缀的响应"""
        response = "润色后：这是润色后的文本"
        cleaned = self.polisher._clean_response(response)
        self.assertEqual(cleaned, "这是润色后的文本")

    def test_clean_response_with_code_block(self):
        """测试清理带代码块的响应"""
        response = "```\n这是文本\n```"
        cleaned = self.polisher._clean_response(response)
        self.assertEqual(cleaned, "这是文本")

    def test_clean_response_normal(self):
        """测试清理正常响应"""
        response = "这是正常的文本"
        cleaned = self.polisher._clean_response(response)
        self.assertEqual(cleaned, "这是正常的文本")


def run_tests():
    """运行所有测试"""
    loader = unittest.TestLoader()
    suite = unittest.TestSuite()

    suite.addTests(loader.loadTestsFromTestCase(TestLLMClient))
    suite.addTests(loader.loadTestsFromTestCase(TestPaperPolisherInit))
    suite.addTests(loader.loadTestsFromTestCase(TestPolishParagraph))
    suite.addTests(loader.loadTestsFromTestCase(TestTrackChanges))
    suite.addTests(loader.loadTestsFromTestCase(TestFootnoteProtection))
    suite.addTests(loader.loadTestsFromTestCase(TestDocumentProcessing))
    suite.addTests(loader.loadTestsFromTestCase(TestResponseCleaning))

    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)

    return result


if __name__ == '__main__':
    print("="*70)
    print("论文润色模块单元测试和集成测试")
    print("="*70)
    print()

    result = run_tests()

    print()
    print("="*70)
    print(f"测试完成: {result.testsRun} 个测试")
    print(f"成功: {result.testsRun - len(result.failures) - len(result.errors)}")
    print(f"失败: {len(result.failures)}")
    print(f"错误: {len(result.errors)}")
    print("="*70)
