import tempfile
import unittest
from pathlib import Path

try:
    from docx import Document
except ModuleNotFoundError as exc:
    raise unittest.SkipTest("python-docx is not installed") from exc

from modules.doc_processor import DocProcessor


class TestDocProcessorPackage(unittest.TestCase):
    def test_extract_document_package_adds_workflow_metadata(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "sample.docx"
            doc = Document()
            doc.core_properties.title = "Sample"
            doc.add_heading("Introduction", level=1)
            doc.add_paragraph("This is a short paragraph.")
            doc.add_heading("Argument", level=2)
            doc.add_paragraph("This paragraph carries the main argument.")
            doc.save(docx_path)

            package = DocProcessor().extract_document_package(str(docx_path))

            self.assertEqual(package["backend"], "script")
            self.assertEqual(package["provider"], "python-docx")
            self.assertFalse(package["needs_review"])
            self.assertGreaterEqual(package["confidence"], 0.8)
            self.assertEqual(len(package["section_tree"]), 1)
            self.assertEqual(package["section_tree"][0]["title"], "Introduction")
            self.assertEqual(package["summary"]["paragraph_count"], 4)
            self.assertTrue(package["plain_text"])

    def test_extract_document_package_from_bytes_preserves_artifact_name(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            buffer_path = Path(temp_dir) / "bytes.docx"
            doc = Document()
            doc.add_paragraph("No heading here.")
            doc.save(buffer_path)
            payload = buffer_path.read_bytes()

        package = DocProcessor().extract_document_package_from_bytes(payload, source_name="upload.docx")

        self.assertTrue(package["needs_review"])
        self.assertIn("no_heading_structure", package["quality_flags"])
        self.assertEqual(package["artifacts"][0]["path"], "upload.docx")


if __name__ == "__main__":
    unittest.main()
