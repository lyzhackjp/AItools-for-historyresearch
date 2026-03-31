import os
import sys
from pathlib import Path

TEST_DIR = Path(__file__).parent
PROJECT_ROOT = TEST_DIR.parent
sys.path.insert(0, str(PROJECT_ROOT))

from config.api_key_manager import load_all_api_keys
load_all_api_keys()

print("="*70)
print("论文润色模块快速测试")
print("="*70)

print("\n[1] 测试模块导入...")
try:
    from modules.paper_polisher import PaperPolisher, create_paper_polisher
    from modules.llm_client import create_llm_client
    print("  ✓ 模块导入成功")
except Exception as e:
    print(f"  ✗ 模块导入失败: {e}")
    sys.exit(1)

print("\n[2] 测试LLM客户端创建...")
try:
    client = create_llm_client({'provider': 'dashscope'})
    print(f"  ✓ LLM客户端创建成功 (provider: {client.provider})")
except Exception as e:
    print(f"  ✗ LLM客户端创建失败: {e}")

print("\n[3] 测试润色器创建...")
try:
    polisher = create_paper_polisher('qwen')
    print(f"  ✓ 润色器创建成功 (api_provider: {polisher.api_provider})")
except Exception as e:
    print(f"  ✗ 润色器创建失败: {e}")

print("\n[4] 测试修订模式功能...")
try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc = Document()
    para = doc.add_paragraph("原始文本")
    
    polisher._apply_track_changes(para, "原始文本", "修改后文本")
    
    del_elem = para._p.find(qn('w:del'))
    ins_elem = para._p.find(qn('w:ins'))
    
    if del_elem is not None and ins_elem is not None:
        print("  ✓ 修订标记创建成功")
    else:
        print("  ✗ 修订标记创建失败")
except Exception as e:
    print(f"  ✗ 修订模式测试失败: {e}")

print("\n[5] 测试脚注引用保护...")
try:
    doc = Document()
    para = doc.add_paragraph()
    
    fn_run = OxmlElement('w:r')
    fn_ref = OxmlElement('w:footnoteReference')
    fn_ref.set(qn('w:id'), '1')
    fn_run.append(fn_ref)
    para._p.append(fn_run)
    
    polisher._apply_track_changes(para, "原始文本", "修改后文本")
    
    fn_refs = para._p.findall('.//w:footnoteReference', 
                              {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    if len(fn_refs) == 1:
        print("  ✓ 脚注引用保护成功")
    else:
        print(f"  ✗ 脚注引用保护失败 (找到 {len(fn_refs)} 个)")
except Exception as e:
    print(f"  ✗ 脚注引用保护测试失败: {e}")

print("\n[6] 测试响应清理...")
try:
    test_cases = [
        ("润色后：这是文本", "这是文本"),
        ("修改后：这是文本", "这是文本"),
        ("这是正常文本", "这是正常文本"),
    ]
    
    all_passed = True
    for input_text, expected in test_cases:
        result = polisher._clean_response(input_text)
        if result != expected:
            all_passed = False
            print(f"  ✗ 响应清理失败: '{input_text}' -> '{result}' (期望: '{expected}')")
    
    if all_passed:
        print("  ✓ 响应清理测试通过")
except Exception as e:
    print(f"  ✗ 响应清理测试失败: {e}")

print("\n" + "="*70)
print("测试完成")
print("="*70)
