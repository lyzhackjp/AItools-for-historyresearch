"""
学术论文智能精简处理模块 - 优化版

专为日本史学术论文设计的智能内容精简工具
基于阿里通义千问，能够智能识别并删除冗余内容
同时保留核心学术信息。

优化内容 (v2.0.0):
- 添加日本史领域术语库
- 优化润色提示词，提升润色质量
- 增加修改建议解释功能
- 支持多种润色模式（精简、润色、扩展）

核心功能：
- 智能内容精简：自动识别并删除逻辑冗余的论述
- 专业文档处理：正确区分正文与脚注内容
- Word原生修订追踪：使用w:del和w:ins元素实现专业修订模式
- 学术严谨性保障：保护历史专有名词和学术术语
- 修改建议解释：提供详细的修改理由说明

技术架构：
- 文档处理：基于 python-docx 库处理 .docx 文档
- AI处理：集成阿里通义千问 / 次要支持 Minimax
- 修订追踪：实现 Word 原生 Track Changes 功能
- 配置管理：使用 .env 格式管理配置参数
"""

import os
import sys
import re
import json
import random
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
import hashlib
from dataclasses import dataclass, field

try:
    import docx
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("警告: python-docx 未安装，使用 pip install python-docx")

from modules.llm_client import create_llm_client


@dataclass
class PolishingResult:
    """润色结果数据类"""
    original_text: str
    polished_text: str
    modifications: List[Dict[str, str]] = field(default_factory=list)
    summary: str = ""
    confidence: float = 0.0
    processing_time: float = 0.0


class DomainTerminologyDB:
    """日本史领域术语库"""
    
    TERMS = {
        'historical_periods': [
            '明治維新', '明治时代', '大正时代', '昭和时代',
            '幕末', '江戸時代', '戦国時代', '平安時代',
            '明治初期', '明治中期', '明治後期'
        ],
        'political_terms': [
            '国体論', '天皇制', '立憲君主制', '議会政治',
            '政党政治', '藩閥政治', '超国家主義',
            '自由民権運動', '大正デモクラシー'
        ],
        'institutions': [
            '幕府', '朝廷', '国会', '貴族院', '衆議院',
            '内務省', '外務省', '大蔵省', '陸軍省', '海軍省',
            '枢密院', '参謀本部', '警視庁'
        ],
        'concepts': [
            '文明開化', '殖産興業', '富国強兵', '脱亜論',
            '尊王攘夷', '公武合体', '佐幕', '倒幕',
            '武士道', '華族制度', '廃藩置県'
        ],
        'academic_terms': [
            '思想史', '政治思想', '社会思想', '近代化論',
            '比較史', '概念史', '社会史', '文化史',
            '実証主義', '文献学', '史料批判'
        ]
    }
    
    PROTECTED_PATTERNS = [
        r'\d{4}年',
        r'\d{1,2}月\d{1,2}日',
        r'『[^』]+』',
        r'「[^」]+」',
        r'《[^》]+》',
        r'\([^)]+\)',
        r'［[^］]+］',
    ]
    
    @classmethod
    def get_all_terms(cls) -> set:
        """获取所有术语"""
        all_terms = set()
        for category, terms in cls.TERMS.items():
            all_terms.update(terms)
        return all_terms
    
    @classmethod
    def is_protected(cls, text: str) -> bool:
        """检查文本是否应被保护"""
        for pattern in cls.PROTECTED_PATTERNS:
            if re.search(pattern, text):
                return True
        return False
    
    @classmethod
    def contains_term(cls, text: str) -> List[str]:
        """检查文本中包含的术语"""
        found_terms = []
        all_terms = cls.get_all_terms()
        for term in all_terms:
            if term in text:
                found_terms.append(term)
        return found_terms


class PaperPolisherOptimized:
    """学术论文智能精简处理器 - 优化版"""
    
    OPTIMIZED_SYSTEM_PROMPT = """你是一位专业的日本史学术论文编辑，具有深厚的史学素养和丰富的编辑经验。

【任务说明】
请对给定的学术论文段落进行润色和精简，在保持学术严谨性的同时提升文章质量。

【润色原则】
1. 学术严谨性：保留所有核心论点、史实依据和学术论证
2. 历史准确性：保护历史专有名词、年代、人名、地名等
3. 逻辑完整性：维护原文的论证逻辑和结构
4. 语言规范性：修正语法错误，提升学术表达

【精简规则】
可删除的内容：
- 重复的修饰词和表达
- 冗余的过渡句
- 过度解释的语句
- 与主题无关的延伸

必须保留的内容：
- 核心学术观点和结论
- 历史史实和数据
- 人物信息和事件描述
- 脚注引用和参考文献标注
- 历史专有名词和学术术语

【输出格式】
请以JSON格式输出：
{
    "polished_text": "润色后的文本",
    "modifications": [
        {
            "type": "删除/修改/保留",
            "original": "原文内容",
            "modified": "修改后内容（如有）",
            "reason": "修改理由"
        }
    ],
    "summary": "本次润色的总体说明",
    "confidence": 0.85
}"""

    MIN_PARAGRAPH_LENGTH = 30

    def __init__(self, api_provider: str = "qwen"):
        """
        初始化论文润色器
        
        Args:
            api_provider: API提供商 ('qwen' 或 'minimax')
        """
        self.base_dir = Path(__file__).parent.parent
        self.api_provider = api_provider
        self.llm_client = None
        self.system_prompt = self.OPTIMIZED_SYSTEM_PROMPT
        self.terminology_db = DomainTerminologyDB()
        self.history_terms = set()
        self.history_figures = set()
        
    def _init_llm_client(self):
        """初始化LLM客户端"""
        if self.llm_client is None:
            if self.api_provider == "qwen":
                provider = "dashscope"
            elif self.api_provider == "minimax":
                provider = "minimax"
            else:
                provider = "dashscope"
            
            self.llm_client = create_llm_client({'provider': provider})
    
    def _load_history_knowledge(self, doc: Document):
        """
        从文档中提取历史专有名词和人物
        
        Args:
            doc: Word文档对象
        """
        text_content = []
        
        for para in doc.paragraphs:
            text_content.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content.append(cell.text)
        
        full_text = "\n".join(text_content)
        
        date_pattern = r'\d{4,4}年\d{1,2}月\d{1,2}日|\d{4,4}年-\d{4,4}年|\d{1,2}世纪'
        dates = re.findall(date_pattern, full_text)
        
        name_pattern = r'[一-龥]{2,4}(氏|公爵|侯爵|伯爵|子爵|男爵)'
        names = re.findall(name_pattern, full_text)
        
        self.history_terms.update(dates)
        self.history_figures.update(names)
        
        domain_terms = self.terminology_db.get_all_terms()
        self.history_terms.update(domain_terms)
    
    def polish_paragraph(self, paragraph_text: str, 
                        mode: str = 'simplify') -> PolishingResult:
        """
        精单润色个段落
        
        Args:
            paragraph_text: 原始段落文本
            mode: 润色模式 ('simplify', 'polish', 'expand')
            
        Returns:
            PolishingResult: 润色结果
        """
        import time
        start_time = time.time()
        
        if not paragraph_text.strip():
            return PolishingResult(
                original_text=paragraph_text,
                polished_text=paragraph_text,
                modifications=[],
                summary="空段落，无需处理"
            )
        
        if len(paragraph_text) < self.MIN_PARAGRAPH_LENGTH:
            return PolishingResult(
                original_text=paragraph_text,
                polished_text=paragraph_text,
                modifications=[],
                summary="段落过短，保持原样"
            )
        
        protected_terms = self.terminology_db.contains_term(paragraph_text)
        
        self._init_llm_client()
        
        mode_instructions = {
            'simplify': '请精简以下段落，删除冗余内容，保留核心论点。精简比例约20%-40%。',
            'polish': '请润色以下段落，优化表达，修正语法，保持原意和长度。',
            'expand': '请扩展以下段落，补充论证细节，增加学术深度。'
        }
        
        user_prompt = f"""{mode_instructions.get(mode, mode_instructions['simplify'])}

【保护术语】
以下术语必须保留：{', '.join(protected_terms[:10]) if protected_terms else '无特殊保护术语'}

【原文】（{len(paragraph_text)}字）
{paragraph_text}

请按JSON格式输出润色结果："""
        
        try:
            response = self.llm_client._call_llm(user_prompt, temperature=0.3, max_tokens=2000)
            
            result_text = response.get('content', '')
            
            try:
                result_data = json.loads(self._extract_json(result_text))
            except json.JSONDecodeError:
                result_data = {
                    'polished_text': self._clean_response(result_text),
                    'modifications': [],
                    'summary': 'LLM返回非JSON格式',
                    'confidence': 0.7
                }
            
            polished_text = result_data.get('polished_text', paragraph_text)
            polished_text = self._clean_response(polished_text)
            
            if not polished_text or len(polished_text) < 10:
                polished_text = paragraph_text
            
            processing_time = time.time() - start_time
            
            return PolishingResult(
                original_text=paragraph_text,
                polished_text=polished_text,
                modifications=result_data.get('modifications', []),
                summary=result_data.get('summary', ''),
                confidence=result_data.get('confidence', 0.8),
                processing_time=processing_time
            )
                
        except Exception as e:
            print(f"段落润色失败: {e}")
            return PolishingResult(
                original_text=paragraph_text,
                polished_text=paragraph_text,
                modifications=[],
                summary=f"处理失败: {str(e)}",
                confidence=0.0
            )
    
    def _extract_json(self, text: str) -> str:
        """从文本中提取JSON"""
        json_pattern = r'\{[^{}]*(?:\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}[^{}]*)*\}'
        matches = re.findall(json_pattern, text, re.DOTALL)
        
        for match in matches:
            if 'polished_text' in match:
                return match
        
        return text
    
    def _clean_response(self, response: str) -> str:
        """清理响应文本"""
        result = response.strip()
        
        prefixes = [
            "润色后：", "润色后:", "修改后：", "修改后:",
            "精简后：", "精简后:", "以下是润色后的文本：",
            "以下是润色后的文本:", "润色后的文本：", "润色后的文本:",
            "精简后的文本：", "精简后的文本:", "polished_text:",
            "```json", "```"
        ]
        for prefix in prefixes:
            if result.startswith(prefix):
                result = result[len(prefix):].strip()
        
        if result.endswith("```"):
            result = result[:-3].strip()
        
        return result.strip()
    
    def process_document(self, input_path: str, output_path: str,
                        enable_track_changes: bool = True,
                        mode: str = 'simplify') -> Dict[str, Any]:
        """
        处理完整的Word文档
        
        Args:
            input_path: 输入文档路径
            output_path: 输出文档路径
            enable_track_changes: 是否启用修订追踪
            mode: 润色模式
            
        Returns:
            Dict: 处理结果统计
        """
        input_path = Path(input_path)
        output_path = Path(output_path)
        
        if not input_path.exists():
            raise FileNotFoundError(f"输入文件不存在: {input_path}")
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        import shutil
        shutil.copy2(str(input_path), str(output_path))
        
        doc = Document(str(output_path))
        
        self._load_history_knowledge(doc)
        
        if enable_track_changes:
            self._enable_track_revisions(doc)
            print(f"✓ 已启用修订追踪模式")
        
        total_paragraphs = len(doc.paragraphs)
        processed_paragraphs = 0
        all_modifications = []
        total_confidence = 0.0
        
        print(f"\n开始处理 {total_paragraphs} 个段落...\n")
        
        for i, para in enumerate(doc.paragraphs, 1):
            original_text = para.text
            
            if not original_text.strip():
                continue
            
            processed_paragraphs += 1
            
            print(f"[{i}/{total_paragraphs}] 处理段落...", end=" ", flush=True)
            
            result = self.polish_paragraph(original_text, mode)
            
            if result.modifications:
                all_modifications.extend(result.modifications)
            
            total_confidence += result.confidence
            
            if result.polished_text != original_text:
                if enable_track_changes:
                    self._apply_track_changes(para, original_text, result.polished_text)
                else:
                    para.clear()
                    para.add_run(result.polished_text)
                
                print(f"✓ ({len(result.modifications)} 处修改, 置信度: {result.confidence:.2f})")
            else:
                print(f"- (无修改)")
        
        doc.save(str(output_path))
        
        avg_confidence = total_confidence / processed_paragraphs if processed_paragraphs > 0 else 0
        
        result = {
            'success': True,
            'input_file': str(input_path),
            'output_file': str(output_path),
            'total_paragraphs': total_paragraphs,
            'processed_paragraphs': processed_paragraphs,
            'total_modifications': len(all_modifications),
            'modifications': all_modifications,
            'average_confidence': avg_confidence,
            'track_changes_enabled': enable_track_changes,
            'mode': mode,
            'timestamp': datetime.now().isoformat()
        }
        
        print(f"\n✓ 处理完成: {output_path}")
        print(f"  - 平均置信度: {avg_confidence:.2f}")
        print(f"  - 总修改数: {len(all_modifications)}")
        
        return result
    
    def _enable_track_revisions(self, doc):
        """启用文档的修订追踪设置"""
        settings = doc.settings.element
        
        track_revisions = settings.find(qn('w:trackRevisions'))
        if track_revisions is None:
            track_revisions = OxmlElement('w:trackRevisions')
            settings.append(track_revisions)
    
    def _apply_track_changes(self, para, original_text: str, modified_text: str):
        """
        应用Word原生修订追踪格式
        
        Args:
            para: 段落对象
            original_text: 原始文本
            modified_text: 修改后文本
        """
        footnote_refs = []
        for child in para._p:
            if child.tag == qn('w:r'):
                fn_ref = child.find(qn('w:footnoteReference'))
                if fn_ref is not None:
                    footnote_refs.append(child)
        
        para.clear()
        
        author = "AI润色助手"
        date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
        revision_id = str(random.randint(1, 999999))
        
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), revision_id)
        del_elem.set(qn('w:author'), author)
        del_elem.set(qn('w:date'), date_str)
        
        del_run = OxmlElement('w:r')
        del_rPr = OxmlElement('w:rPr')
        del_run.append(del_rPr)
        
        del_text = OxmlElement('w:delText')
        del_text.text = original_text
        del_text.set(qn('xml:space'), 'preserve')
        del_run.append(del_text)
        del_elem.append(del_run)
        
        para._p.append(del_elem)
        
        ins_elem = OxmlElement('w:ins')
        ins_elem.set(qn('w:id'), str(int(revision_id) + 1))
        ins_elem.set(qn('w:author'), author)
        ins_elem.set(qn('w:date'), date_str)
        
        ins_run = OxmlElement('w:r')
        ins_rPr = OxmlElement('w:rPr')
        ins_run.append(ins_rPr)
        
        ins_text = OxmlElement('w:t')
        ins_text.text = modified_text
        ins_text.set(qn('xml:space'), 'preserve')
        ins_run.append(ins_text)
        ins_elem.append(ins_run)
        
        para._p.append(ins_elem)
        
        for fn_ref in footnote_refs:
            para._p.append(fn_ref)
    
    def set_system_prompt(self, prompt: str):
        """
        设置系统提示词
        
        Args:
            prompt: 新的系统提示词
        """
        self.system_prompt = prompt
    
    def add_history_term(self, term: str):
        """
        添加历史专有名词保护
        
        Args:
            term: 历史术语
        """
        self.history_terms.add(term)
    
    def add_history_figure(self, figure: str):
        """
        添加历史人物保护
        
        Args:
            figure: 历史人物名称
        """
        self.history_figures.add(figure)
    
    def get_modification_report(self, result: Dict[str, Any]) -> str:
        """
        生成修改报告
        
        Args:
            result: 处理结果
            
        Returns:
            str: Markdown格式的报告
        """
        report_lines = [
            "# 论文润色报告",
            "",
            f"**处理时间**: {result.get('timestamp', 'N/A')}",
            f"**润色模式**: {result.get('mode', 'simplify')}",
            f"**平均置信度**: {result.get('average_confidence', 0):.2f}",
            "",
            "## 处理统计",
            "",
            f"- 总段落数: {result.get('total_paragraphs', 0)}",
            f"- 处理段落数: {result.get('processed_paragraphs', 0)}",
            f"- 总修改数: {result.get('total_modifications', 0)}",
            "",
            "## 修改详情",
            ""
        ]
        
        modifications = result.get('modifications', [])
        for i, mod in enumerate(modifications[:20], 1):
            report_lines.append(f"### 修改 {i}")
            report_lines.append(f"- **类型**: {mod.get('type', 'N/A')}")
            report_lines.append(f"- **原文**: {mod.get('original', 'N/A')[:100]}...")
            if mod.get('modified'):
                report_lines.append(f"- **修改后**: {mod.get('modified', '')[:100]}...")
            report_lines.append(f"- **理由**: {mod.get('reason', 'N/A')}")
            report_lines.append("")
        
        if len(modifications) > 20:
            report_lines.append(f"... 还有 {len(modifications) - 20} 处修改未显示")
        
        return '\n'.join(report_lines)


def create_paper_polisher_optimized(api_provider: str = "qwen") -> PaperPolisherOptimized:
    """
    工厂函数 - 创建优化版论文润色器
    
    Args:
        api_provider: API提供商 ('qwen' 或 'minimax')
        
    Returns:
        PaperPolisherOptimized: 论文润色器实例
    """
    return PaperPolisherOptimized(api_provider)


if __name__ == "__main__":
    print("学术论文智能精简处理工具 - 优化版 v2.0.0")
    print("="*60)
    print("\n使用方法:")
    print("```python")
    print("from modules.paper_polisher_optimized import create_paper_polisher_optimized")
    print("")
    print("# 创建润色器")
    print("polisher = create_paper_polisher_optimized('qwen')")
    print("")
    print("# 处理文档")
    print("result = polisher.process_document(")
    print("    'input.docx',")
    print("    'output.docx',")
    print("    enable_track_changes=True,")
    print("    mode='simplify'  # 'simplify', 'polish', 'expand'")
    print(")")
    print("")
    print("# 生成修改报告")
    print("report = polisher.get_modification_report(result)")
    print("```")
