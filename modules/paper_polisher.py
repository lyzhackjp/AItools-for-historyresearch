"""
Academic paper polishing facade.

This module keeps the old public surface area while routing the actual text
polishing work through the unified task layer.
"""

from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from modules.task_manager import TaskManager

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - optional dependency
    Document = None
    OxmlElement = None
    qn = None
    DOCX_AVAILABLE = False


class PaperPolisher:
    """Conservative writing polisher for academic prose."""

    DEFAULT_SYSTEM_PROMPT = (
        "Polish academic writing conservatively while preserving claims, evidence, "
        "citations, and footnotes."
    )
    MIN_PARAGRAPH_LENGTH = 30

    def __init__(
        self,
        api_provider: str = "qwen",
        test_mode: bool = True,
        backend: Optional[str] = None,
        model: Optional[str] = None,
    ):
        self.api_provider = api_provider
        self.test_mode = test_mode
        self.backend = backend
        self.model = model
        self.llm_client = None
        self.system_prompt = self.DEFAULT_SYSTEM_PROMPT
        self._task_manager: Optional[TaskManager] = None

    def _get_task_manager(self) -> TaskManager:
        if self._task_manager is None:
            self._task_manager = TaskManager(mode="api", provider=self.api_provider)
        return self._task_manager

    def _init_llm_client(self) -> TaskManager:
        """Compatibility shim for older callers."""

        self.llm_client = self._get_task_manager()
        return self.llm_client

    def get_capabilities(self) -> Dict[str, Any]:
        task_options = self._get_task_manager().get_task_options("paper_polish")
        task_options.update(
            {
                "module": "PaperPolisher",
                "task": "paper_polish",
                "output_type": "paper_polish",
                "test_mode": self.test_mode,
                "legacy_methods": ["polish_paragraph", "polish_text", "process_document"],
                "fallback_order": ["llm_api", "local_llm", "script", "skill", "mcp"],
                "quality_signals": [
                    "empty_input",
                    "short_input_skipped",
                    "empty_output",
                    "fallback_backend",
                    "low_confidence",
                    "suspicious_length_change",
                ],
            }
        )
        return task_options

    def polish_paragraph(
        self,
        paragraph_text: str,
        language: str = "zh",
        **kwargs: Any,
    ) -> Tuple[str, List[Dict[str, str]]]:
        """Polish a single paragraph and return `(text, revision_notes)`."""

        if not paragraph_text.strip():
            return paragraph_text, []
        if len(paragraph_text.strip()) < self.MIN_PARAGRAPH_LENGTH:
            return paragraph_text, []

        result = self._execute_polish(paragraph_text, language=language, **kwargs)
        polished_text = result.get("polished_text") or paragraph_text
        revision_notes = self._normalize_revision_notes(
            result.get("revision_notes", []),
            original_text=paragraph_text,
            polished_text=polished_text,
        )
        return polished_text, revision_notes

    def polish_text(
        self,
        text: str,
        language: str = "zh",
        **kwargs: Any,
    ) -> Dict[str, Any]:
        """Polish multi-paragraph text and keep execution metadata."""

        paragraphs = self._split_paragraphs(text)
        if not paragraphs:
            return {
                "polished_text": text,
                "revision_notes": [],
                "backend": self.backend or ("script" if self.test_mode else None),
                "provider": self.api_provider,
                "model": self.model,
                "confidence": 0.0,
                "needs_review": False,
            }

        polished_paragraphs: List[str] = []
        revision_notes: List[Dict[str, str]] = []
        backends_used: List[str] = []
        providers_used: List[str] = []
        models_used: List[str] = []
        needs_review = False
        confidence_values: List[float] = []

        for paragraph in paragraphs:
            if len(paragraph.strip()) < self.MIN_PARAGRAPH_LENGTH:
                polished_paragraphs.append(paragraph)
                continue

            result = self._execute_polish(paragraph, language=language, **kwargs)
            polished_text = result.get("polished_text") or paragraph
            polished_paragraphs.append(polished_text)
            revision_notes.extend(
                self._normalize_revision_notes(
                    result.get("revision_notes", []),
                    original_text=paragraph,
                    polished_text=polished_text,
                )
            )
            if result.get("backend"):
                backends_used.append(str(result["backend"]))
            if result.get("provider"):
                providers_used.append(str(result["provider"]))
            if result.get("model"):
                models_used.append(str(result["model"]))
            if isinstance(result.get("confidence"), (int, float)):
                confidence_values.append(float(result["confidence"]))
            needs_review = needs_review or bool(result.get("needs_review"))

        if not polished_paragraphs:
            polished_paragraphs = paragraphs

        return {
            "polished_text": "\n\n".join(polished_paragraphs),
            "revision_notes": revision_notes,
            "backend": backends_used[0] if backends_used else self.backend or ("script" if self.test_mode else None),
            "provider": providers_used[0] if providers_used else self.api_provider,
            "model": models_used[0] if models_used else self.model,
            "confidence": round(sum(confidence_values) / len(confidence_values), 2) if confidence_values else 0.55,
            "needs_review": needs_review,
            "backend_chain": backends_used,
        }

    def polish_paragraph_package(
        self,
        paragraph_text: str,
        language: str = "zh",
        **kwargs: Any,
    ) -> Dict[str, Any]:
        """Polish one paragraph and return a `paper_polish` envelope."""

        result = self.polish_text(paragraph_text, language=language, **kwargs)
        package = self._build_polish_package(
            original_text=paragraph_text,
            polish_result=result,
            language=language,
            scope="paragraph",
        )
        if paragraph_text.strip() and len(paragraph_text.strip()) < self.MIN_PARAGRAPH_LENGTH:
            package["quality_flags"] = sorted(set(package["quality_flags"] + ["short_input_skipped"]))
            package["needs_review"] = bool(package["quality_flags"])
        return package

    def polish_text_package(
        self,
        text: str,
        language: str = "zh",
        **kwargs: Any,
    ) -> Dict[str, Any]:
        """Polish text and return a workflow-friendly `paper_polish` envelope."""

        result = self.polish_text(text, language=language, **kwargs)
        return self._build_polish_package(
            original_text=text,
            polish_result=result,
            language=language,
            scope="text",
        )

    def process_document(
        self,
        input_path: str,
        output_path: str,
        enable_track_changes: bool = True,
        language: str = "zh",
        **kwargs: Any,
    ) -> Dict[str, Any]:
        """Polish a `.docx` file while preserving footnote references."""

        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is required for process_document")

        source = Path(input_path)
        target = Path(output_path)
        if not source.exists():
            raise FileNotFoundError(f"Input file does not exist: {source}")

        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, target)

        doc = Document(str(target))
        if enable_track_changes:
            self._enable_track_revisions(doc)

        processed_paragraphs = 0
        modified_paragraphs = 0
        all_notes: List[Dict[str, str]] = []
        backends_used: List[str] = []

        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if not original_text.strip():
                continue
            processed_paragraphs += 1
            result = self._execute_polish(original_text, language=language, **kwargs)
            polished_text = result.get("polished_text") or original_text
            if polished_text != original_text:
                modified_paragraphs += 1
                if enable_track_changes:
                    self._apply_track_changes(paragraph, original_text, polished_text)
                else:
                    paragraph.clear()
                    paragraph.add_run(polished_text)
            all_notes.extend(
                self._normalize_revision_notes(
                    result.get("revision_notes", []),
                    original_text=original_text,
                    polished_text=polished_text,
                )
            )
            if result.get("backend"):
                backends_used.append(str(result["backend"]))

        doc.save(str(target))
        return {
            "success": True,
            "input_file": str(source),
            "output_file": str(target),
            "total_paragraphs": len(doc.paragraphs),
            "processed_paragraphs": processed_paragraphs,
            "modified_paragraphs": modified_paragraphs,
            "total_deletions": len(all_notes),
            "revision_notes": all_notes,
            "backend_chain": backends_used,
        }

    def process_document_package(
        self,
        input_path: str,
        output_path: str,
        enable_track_changes: bool = True,
        language: str = "zh",
        **kwargs: Any,
    ) -> Dict[str, Any]:
        """Process a DOCX and wrap the result as a `paper_polish_document` package."""

        result = self.process_document(
            input_path,
            output_path,
            enable_track_changes=enable_track_changes,
            language=language,
            **kwargs,
        )
        flags = []
        if not result.get("success"):
            flags.append("document_polish_failed")
        if not result.get("modified_paragraphs"):
            flags.append("no_paragraphs_modified")
        return {
            "type": "paper_polish_document",
            "schema_version": "2026-04-25",
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "source_path": result.get("input_file"),
            "output_path": result.get("output_file"),
            "backend": result.get("backend_chain", [None])[0] if result.get("backend_chain") else "script",
            "provider": self.api_provider,
            "model": self.model,
            "confidence": 0.7 if result.get("success") else 0.0,
            "needs_review": bool(flags),
            "quality_flags": flags,
            "statistics": {
                "total_paragraphs": result.get("total_paragraphs", 0),
                "processed_paragraphs": result.get("processed_paragraphs", 0),
                "modified_paragraphs": result.get("modified_paragraphs", 0),
                "revision_note_count": len(result.get("revision_notes", [])),
            },
            "revision_notes": result.get("revision_notes", []),
            "artifacts": [
                {
                    "kind": "docx",
                    "path": result.get("output_file"),
                }
            ] if result.get("output_file") else [],
            "capabilities": self.get_capabilities(),
            "raw_result": result,
        }

    def _execute_polish(self, text: str, language: str = "zh", **kwargs: Any) -> Dict[str, Any]:
        manager = self._get_task_manager()
        backend = kwargs.get("backend", self.backend)
        fallback_backends = kwargs.get("fallback_backends")
        if fallback_backends is None:
            fallback_backends = ["local_llm", "script"]
        if self.test_mode and not backend:
            backend = "script"
            fallback_backends = []

        response = manager.paper_polish(
            text=text,
            language=language,
            provider=kwargs.get("provider", self.api_provider),
            model=kwargs.get("model", self.model),
            backend=backend,
            fallback_backends=list(fallback_backends),
            temperature=kwargs.get("temperature", 0.2),
            max_tokens=kwargs.get("max_tokens", 2500),
        )

        payload = response.get("data", {}) if response.get("success") else {}
        polished_text = payload.get("polished_text") or text
        revision_notes = payload.get("revision_notes", [])
        if not response.get("success"):
            revision_notes = revision_notes or ["polish fallback returned original text"]

        return {
            "polished_text": self._clean_response(polished_text),
            "revision_notes": revision_notes,
            "backend": response.get("backend") or response.get("metadata", {}).get("backend"),
            "provider": response.get("metadata", {}).get("provider", self.api_provider),
            "model": response.get("metadata", {}).get("model", self.model),
            "confidence": payload.get("confidence", 0.7 if polished_text != text else 0.55),
            "needs_review": bool(payload.get("needs_review", False)),
        }

    def _build_polish_package(
        self,
        *,
        original_text: str,
        polish_result: Dict[str, Any],
        language: str,
        scope: str,
    ) -> Dict[str, Any]:
        flags = self._package_quality_flags(original_text, polish_result)
        confidence = self._package_confidence(polish_result, flags)
        return {
            "type": "paper_polish",
            "schema_version": "2026-04-25",
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "scope": scope,
            "language": language,
            "original_text": original_text,
            "polished_text": polish_result.get("polished_text", original_text),
            "revision_notes": polish_result.get("revision_notes", []),
            "backend": polish_result.get("backend") or ("script" if self.test_mode else self.backend),
            "provider": polish_result.get("provider", self.api_provider),
            "model": polish_result.get("model", self.model),
            "confidence": confidence,
            "needs_review": bool(flags) or bool(polish_result.get("needs_review")),
            "quality_flags": flags,
            "statistics": {
                "original_chars": len(original_text or ""),
                "polished_chars": len(polish_result.get("polished_text", "") or ""),
                "revision_note_count": len(polish_result.get("revision_notes", [])),
                "paragraph_count": len(self._split_paragraphs(original_text)),
            },
            "backend_chain": polish_result.get("backend_chain", []),
            "capabilities": self.get_capabilities(),
        }

    def _package_quality_flags(self, original_text: str, polish_result: Dict[str, Any]) -> List[str]:
        flags = []
        original = original_text or ""
        polished = polish_result.get("polished_text", "") or ""
        if not original.strip():
            flags.append("empty_input")
        if original.strip() and not polished.strip():
            flags.append("empty_output")
        if polish_result.get("backend") in {"script", "fallback"} and not self.test_mode:
            flags.append("fallback_backend")
        if polish_result.get("needs_review"):
            flags.append("backend_review_requested")
        confidence = polish_result.get("confidence")
        if isinstance(confidence, (int, float)) and confidence < 0.6:
            flags.append("low_confidence")
        if original.strip() and polished.strip():
            ratio = len(polished) / max(len(original), 1)
            if ratio < 0.3 or ratio > 1.8:
                flags.append("suspicious_length_change")
        return sorted(set(flags))

    def _package_confidence(self, polish_result: Dict[str, Any], flags: List[str]) -> float:
        confidence = polish_result.get("confidence")
        if not isinstance(confidence, (int, float)):
            confidence = 0.65
        if "fallback_backend" in flags:
            confidence -= 0.1
        if "empty_output" in flags:
            confidence -= 0.35
        if "suspicious_length_change" in flags:
            confidence -= 0.15
        if "low_confidence" in flags:
            confidence = min(confidence, 0.55)
        return round(max(0.0, min(1.0, float(confidence))), 3)

    def _normalize_revision_notes(
        self,
        notes: List[Any],
        *,
        original_text: str,
        polished_text: str,
    ) -> List[Dict[str, str]]:
        normalized: List[Dict[str, str]] = []
        if isinstance(notes, list):
            for note in notes:
                if isinstance(note, dict):
                    normalized.append(
                        {
                            "text": str(note.get("text") or note.get("reason") or ""),
                            "reason": str(note.get("reason") or note.get("text") or ""),
                        }
                    )
                elif isinstance(note, str):
                    normalized.append({"text": note, "reason": note})
        if not normalized and polished_text != original_text:
            normalized.append(
                {
                    "text": "conservative rewrite applied",
                    "reason": f"reduced text length from {len(original_text)} to {len(polished_text)} characters",
                }
            )
        return [item for item in normalized if item.get("text") or item.get("reason")]

    def _split_paragraphs(self, text: str) -> List[str]:
        paragraphs = [part.strip() for part in text.split("\n\n") if part.strip()]
        if paragraphs:
            return paragraphs
        return [part.strip() for part in text.splitlines() if part.strip()]

    def _clean_response(self, response: str) -> str:
        text = (response or "").strip()
        prefixes = [
            "润色后：",
            "润色后:",
            "修改后：",
            "修改后:",
            "精简后：",
            "精简后:",
            "以下是润色后的文本：",
            "以下是润色后的文本:",
        ]
        for prefix in prefixes:
            if text.startswith(prefix):
                text = text[len(prefix) :].strip()

        if text.startswith("```"):
            lines = text.splitlines()
            if lines and lines[0].startswith("```"):
                lines = lines[1:]
            if lines and lines[-1].startswith("```"):
                lines = lines[:-1]
            text = "\n".join(lines).strip()
        return text

    def _enable_track_revisions(self, doc: Any) -> None:
        track_revisions = doc.settings.element.find(qn("w:trackRevisions"))
        if track_revisions is None:
            track_revisions = OxmlElement("w:trackRevisions")
            doc.settings.element.append(track_revisions)

    def _apply_track_changes(self, paragraph: Any, original_text: str, modified_text: str) -> None:
        footnote_runs = []
        for child in list(paragraph._p):
            if child.find(qn("w:footnoteReference")) is not None:
                footnote_runs.append(child)
            paragraph._p.remove(child)

        if original_text:
            deletion = OxmlElement("w:del")
            deletion.set(qn("w:author"), "Codex")
            deletion.set(qn("w:date"), "2026-04-21T00:00:00Z")
            run = OxmlElement("w:r")
            text = OxmlElement("w:delText")
            text.text = original_text
            run.append(text)
            deletion.append(run)
            paragraph._p.append(deletion)

        if modified_text:
            insertion = OxmlElement("w:ins")
            insertion.set(qn("w:author"), "Codex")
            insertion.set(qn("w:date"), "2026-04-21T00:00:00Z")
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = modified_text
            run.append(text)
            insertion.append(run)
            paragraph._p.append(insertion)

        for footnote_run in footnote_runs:
            paragraph._p.append(footnote_run)


def create_paper_polisher(api_provider: str = "qwen", **kwargs: Any) -> PaperPolisher:
    """Compatibility factory."""

    return PaperPolisher(api_provider=api_provider, **kwargs)
