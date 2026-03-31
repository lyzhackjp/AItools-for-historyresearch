from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import zipfile
from typing import Dict, Any, List, Optional
from lxml import etree


class DocProcessor:
    """Word文档处理模块 - 支持解析与生成.docx格式文档，保留完整格式"""

    def __init__(self):
        """初始化文档处理器"""
        pass

    def extract_text(self, file_path: str) -> dict:
        """
        从Word文档中提取文本内容（保留完整格式信息）

        Args:
            file_path: docx文件路径

        Returns:
            dict: 包含文档结构和文本内容的字典
        """
        doc = Document(file_path)
        return self._parse_document(doc, file_path)

    def extract_text_from_bytes(self, file_bytes: bytes) -> dict:
        """从字节流提取文档内容"""
        doc = Document(io.BytesIO(file_bytes))
        return self._parse_document(doc)

    def _parse_document(self, doc: Document, file_path: str = None) -> dict:
        """
        解析Word文档，提取所有元素

        Args:
            doc: Document对象
            file_path: 文件路径（用于提取脚注/尾注）

        Returns:
            dict: 文档解析结果
        """
        result = {
            'title': '',
            'paragraphs': [],
            'tables': [],
            'styles': {},
            'headers': [],
            'footers': [],
            'footnotes': [],
            'endnotes': [],
            'page_numbers': [],
            'metadata': {}
        }

        result['title'] = self._extract_title(doc)
        result['metadata'] = self._extract_metadata(doc)

        for section in doc.sections:
            result['headers'].extend(self._extract_header(section))
            result['footers'].extend(self._extract_footer(section))

        if file_path:
            result['footnotes'] = self._extract_footnotes_from_xml(file_path)
            result['endnotes'] = self._extract_endnotes_from_xml(file_path)
        else:
            result['footnotes'] = self._extract_footnotes(doc)
            result['endnotes'] = self._extract_endnotes(doc)

        for para in doc.paragraphs:
            para_info = self._parse_paragraph(para)
            if para_info['text'].strip():
                result['paragraphs'].append(para_info)
                result['styles'][para.style.name] = para_info['text'][:50]

        for table in doc.tables:
            table_data = self._parse_table(table)
            result['tables'].append(table_data)

        return result

    def _extract_footnotes_from_xml(self, file_path: str) -> List[Dict[str, Any]]:
        """
        从footnotes.xml中提取脚注内容

        Args:
            file_path: docx文件路径

        Returns:
            list: 脚注列表
        """
        footnotes = []

        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                try:
                    with z.open('word/footnotes.xml') as f:
                        content = f.read()
                        tree = etree.fromstring(content)

                        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        footnote_elements = tree.findall('.//w:footnote', ns)

                        for fn in footnote_elements:
                            fn_id = fn.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                            fn_type = fn.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'normal')

                            text_parts = []
                            for t in fn.findall('.//w:t', ns):
                                if t.text:
                                    text_parts.append(t.text)

                            text = ''.join(text_parts).strip()

                            if fn_id and text and fn_type not in ('separator', 'continuationSeparator'):
                                footnotes.append({
                                    'id': fn_id,
                                    'type': fn_type,
                                    'text': text
                                })

                except KeyError:
                    pass

        except Exception:
            pass

        return footnotes

    def _extract_endnotes_from_xml(self, file_path: str) -> List[Dict[str, Any]]:
        """
        从endnotes.xml中提取尾注内容

        Args:
            file_path: docx文件路径

        Returns:
            list: 尾注列表
        """
        endnotes = []

        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                try:
                    with z.open('word/endnotes.xml') as f:
                        content = f.read()
                        tree = etree.fromstring(content)

                        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        endnote_elements = tree.findall('.//w:endnote', ns)

                        for en in endnote_elements:
                            en_id = en.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                            en_type = en.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'normal')

                            text_parts = []
                            for t in en.findall('.//w:t', ns):
                                if t.text:
                                    text_parts.append(t.text)

                            text = ''.join(text_parts).strip()

                            if en_id and text and en_type not in ('separator', 'continuationSeparator'):
                                endnotes.append({
                                    'id': en_id,
                                    'type': en_type,
                                    'text': text
                                })

                except KeyError:
                    pass

        except Exception:
            pass

        return endnotes

    def _extract_title(self, doc: Document) -> str:
        """提取文档标题"""
        if doc.core_properties.title:
            return doc.core_properties.title

        for para in doc.paragraphs:
            if para.style.name.startswith('Title') or para.style.name.startswith('Heading 1'):
                return para.text.strip()

        return ''

    def _extract_metadata(self, doc: Document) -> dict:
        """提取文档元数据"""
        core = doc.core_properties
        return {
            'title': core.title,
            'author': core.author,
            'subject': core.subject,
            'keywords': core.keywords,
            'created': str(core.created) if core.created else None,
            'modified': str(core.modified) if core.modified else None
        }

    def _extract_header(self, section) -> List[Dict[str, Any]]:
        """提取页眉内容"""
        headers = []
        header_types = ['primary', 'first', 'even']

        for header_type in header_types:
            try:
                header = section.header
                if header_type == 'primary':
                    header = section.header
                elif header_type == 'first':
                    header = section.first_page_header
                elif header_type == 'even':
                    header = section.even_page_header

                for para in header.paragraphs:
                    if para.text.strip():
                        headers.append({
                            'type': header_type,
                            'text': para.text.strip(),
                            'alignment': str(para.alignment) if para.alignment else 'LEFT'
                        })
            except:
                pass

        return headers

    def _extract_footer(self, section) -> List[Dict[str, Any]]:
        """提取页脚内容"""
        footers = []
        footer_types = ['primary', 'first', 'even']

        for footer_type in footer_types:
            try:
                if footer_type == 'primary':
                    footer = section.footer
                elif footer_type == 'first':
                    footer = section.first_page_footer
                elif footer_type == 'even':
                    footer = section.even_page_footer

                for para in footer.paragraphs:
                    if para.text.strip():
                        footers.append({
                            'type': footer_type,
                            'text': para.text.strip(),
                            'alignment': str(para.alignment) if para.alignment else 'LEFT'
                        })
            except:
                pass

        return footers

    def _extract_footnotes(self, doc: Document) -> List[Dict[str, Any]]:
        """
        提取脚注

        注意: python-docx对脚注的支持有限，此处提取脚注引用标记
        """
        footnotes = []

        for para in doc.paragraphs:
            for run in para.runs:
                if run._element.xpath('.//w:footnoteReference'):
                    footnotes.append({
                        'text': run.text,
                        'reference': 'footnote'
                    })

        return footnotes

    def _extract_endnotes(self, doc: Document) -> List[Dict[str, Any]]:
        """
        提取尾注

        注意: python-docx对尾注的支持有限，此处提取尾注引用标记
        """
        endnotes = []

        for para in doc.paragraphs:
            for run in para.runs:
                if run._element.xpath('.//w:endnoteReference'):
                    endnotes.append({
                        'text': run.text,
                        'reference': 'endnote'
                    })

        return endnotes

    def _parse_paragraph(self, para) -> Dict[str, Any]:
        """
        解析段落，提取详细信息

        Args:
            para: 段落对象

        Returns:
            dict: 段落信息
        """
        para_info = {
            'text': para.text,
            'style': para.style.name if para.style else 'Normal',
            'alignment': str(para.alignment) if para.alignment else 'LEFT',
            'line_spacing': self._get_line_spacing(para),
            'font': self._get_paragraph_font(para),
            'indentation': self._get_indentation(para),
            'runs': self._parse_runs(para)
        }

        return para_info

    def _get_line_spacing(self, para) -> Dict[str, Any]:
        """获取段落行距信息"""
        pf = para.paragraph_format
        spacing_info = {
            'line_spacing_rule': str(pf.line_spacing_rule) if pf.line_spacing_rule else None,
            'line_spacing': pf.line_spacing,
            'space_before': pf.space_before,
            'space_after': pf.space_after
        }
        return spacing_info

    def _get_paragraph_font(self, para) -> Dict[str, Any]:
        """获取段落字体信息"""
        if para.runs:
            first_run = para.runs[0]
            return {
                'name': first_run.font.name,
                'size': first_run.font.size,
                'bold': first_run.font.bold,
                'italic': first_run.font.italic,
                'underline': first_run.font.underline,
                'color': str(first_run.font.color.rgb) if first_run.font.color and first_run.font.color.rgb else None
            }
        return {}

    def _get_indentation(self, para) -> Dict[str, Any]:
        """获取段落缩进信息"""
        pf = para.paragraph_format
        return {
            'left': pf.left_indent,
            'right': pf.right_indent,
            'first_line': pf.first_line_indent
        }

    def _parse_runs(self, para) -> List[Dict[str, Any]]:
        """
        解析文本块（runs），提取每个文本片段的格式

        Args:
            para: 段落对象

        Returns:
            list: 文本块列表
        """
        runs_info = []

        for run in para.runs:
            run_info = {
                'text': run.text,
                'font': {
                    'name': run.font.name,
                    'size': run.font.size,
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'underline': run.font.underline,
                    'color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None
                }
            }
            runs_info.append(run_info)

        return runs_info

    def _parse_table(self, table) -> Dict[str, Any]:
        """
        解析表格

        Args:
            table: 表格对象

        Returns:
            dict: 表格信息
        """
        table_data = {
            'rows': len(table.rows),
            'columns': len(table.columns),
            'data': [],
            'style': table.style.name if table.style else None,
            'alignment': str(table.alignment) if table.alignment else None
        }

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cell_info = {
                    'text': cell_text,
                    'paragraphs': [self._parse_paragraph(p) for p in cell.paragraphs if p.text.strip()]
                }
                row_data.append(cell_info)
            table_data['data'].append(row_data)

        return table_data

    def create_document(self, content: dict, output_path: str, preserve_footnotes: bool = True) -> bool:
        """
        根据处理后的内容创建Word文档（保留格式和脚注）

        Args:
            content: 包含文本内容的字典
            output_path: 输出文件路径
            preserve_footnotes: 是否保留脚注

        Returns:
            bool: 创建是否成功
        """
        doc = Document()

        if content.get('title'):
            title = doc.add_heading(content['title'], level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if content.get('metadata'):
            for key, value in content['metadata'].items():
                if value and hasattr(doc.core_properties, key):
                    try:
                        if key in ['created', 'modified'] and isinstance(value, str):
                            import datetime
                            for fmt in ['%Y-%m-%d %H:%M:%S%z', '%Y-%m-%d %H:%M:%S']:
                                try:
                                    dt = datetime.datetime.strptime(value.replace('+00:00', ''), fmt)
                                    setattr(doc.core_properties, key, dt)
                                    break
                                except:
                                    continue
                            else:
                                continue
                        else:
                            setattr(doc.core_properties, key, value)
                    except:
                        pass

        for para_info in content.get('paragraphs', []):
            para = self._create_paragraph(doc, para_info)

        for table_data in content.get('tables', []):
            if table_data:
                table = self._create_table(doc, table_data)

        for header_info in content.get('headers', []):
            self._create_header(doc, header_info)

        for footer_info in content.get('footers', []):
            self._create_footer(doc, footer_info)

        if preserve_footnotes and content.get('footnotes'):
            self._create_footnotes(doc, content['footnotes'])

        if preserve_footnotes and content.get('endnotes'):
            self._create_endnotes(doc, content['endnotes'])

        doc.save(output_path)
        return True

    def _create_footnotes(self, doc: Document, footnotes: List[Dict[str, Any]]):
        """
        创建脚注

        Args:
            doc: Document对象
            footnotes: 脚注列表
        """
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from lxml import etree

        footnotes_xml = b'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:footnote w:id="0" w:type="separator">
        <w:p>
            <w:r>
                <w:separator/>
            </w:r>
        </w:p>
    </w:footnote>
    <w:footnote w:id="-1" w:type="continuationSeparator">
        <w:p>
            <w:r>
                <w:continuationSeparator/>
            </w:r>
        </w:p>
    </w:footnote>
</w:footnotes>'''

        tree = etree.fromstring(footnotes_xml)
        footnotes_root = tree

        for fn in footnotes:
            fn_id = int(fn.get('id', 1))
            fn_text = fn.get('text', '')

            footnote_elem = etree.SubElement(footnotes_root, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnote')
            footnote_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(fn_id))
            footnote_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'normal')

            p_elem = etree.SubElement(footnote_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
            p_pr = etree.SubElement(p_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            p_style = etree.SubElement(p_pr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
            p_style.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'FootnoteText')

            r_elem1 = etree.SubElement(p_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            r_pr = etree.SubElement(r_elem1, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
            r_style = etree.SubElement(r_pr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rStyle')
            r_style.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'FootnoteReference')
            fn_ref = etree.SubElement(r_elem1, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnoteRef')

            r_elem2 = etree.SubElement(p_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            t_elem = etree.SubElement(r_elem2, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            t_elem.text = fn_text

        footnotes_xml_str = etree.tostring(tree, encoding='unicode')
        footnotes_bytes = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + footnotes_xml_str).encode('utf-8')

        reltype = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
        partname = PackURI('/word/footnotes.xml')
        part = Part(partname, 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml', footnotes_bytes)
        doc.part.package.relate_to(part, reltype)

    def _create_endnotes(self, doc: Document, endnotes: List[Dict[str, Any]]):
        """
        创建尾注

        Args:
            doc: Document对象
            endnotes: 尾注列表
        """
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from lxml import etree

        endnotes_xml = b'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:endnote w:id="0" w:type="separator">
        <w:p>
            <w:r>
                <w:separator/>
            </w:r>
        </w:p>
    </w:endnote>
</w:endnotes>'''

        tree = etree.fromstring(endnotes_xml)
        endnotes_root = tree

        for en in endnotes:
            en_id = int(en.get('id', 1))
            en_text = en.get('text', '')

            endnote_elem = etree.SubElement(endnotes_root, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}endnote')
            endnote_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(en_id))
            endnote_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'normal')

            p_elem = etree.SubElement(endnote_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')

            r_elem1 = etree.SubElement(p_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            en_ref = etree.SubElement(r_elem1, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}endnoteRef')

            r_elem2 = etree.SubElement(p_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            t_elem = etree.SubElement(r_elem2, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            t_elem.text = en_text

        endnotes_xml_str = etree.tostring(tree, encoding='unicode')
        endnotes_bytes = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + endnotes_xml_str).encode('utf-8')

        reltype = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes'
        partname = PackURI('/word/endnotes.xml')
        part = Part(partname, 'application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml', endnotes_bytes)
        doc.part.package.relate_to(part, reltype)

    def _create_paragraph(self, doc: Document, para_info: dict):
        """创建段落并应用格式"""
        para = doc.add_paragraph()

        alignment_map = {
            'CENTER': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'LEFT': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'RIGHT': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'JUSTIFY': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }

        if para_info.get('alignment') in alignment_map:
            para.alignment = alignment_map[para_info['alignment']]

        if para_info.get('line_spacing'):
            spacing = para_info['line_spacing']
            if spacing.get('line_spacing'):
                para.paragraph_format.line_spacing = spacing['line_spacing']
            if spacing.get('space_before'):
                para.paragraph_format.space_before = spacing['space_before']
            if spacing.get('space_after'):
                para.paragraph_format.space_after = spacing['space_after']

        if para_info.get('indentation'):
            indent = para_info['indentation']
            if indent.get('left'):
                para.paragraph_format.left_indent = indent['left']
            if indent.get('right'):
                para.paragraph_format.right_indent = indent['right']
            if indent.get('first_line'):
                para.paragraph_format.first_line_indent = indent['first_line']

        if 'Heading' in para_info.get('style', ''):
            level = 1 if '1' in para_info['style'] else 2 if '2' in para_info['style'] else 3
            run = para.add_run(para_info['text'])
            run.bold = True
            run.font.size = Pt(14 - level)
        else:
            runs_info = para_info.get('runs', [])
            if runs_info:
                for run_info in runs_info:
                    run = para.add_run(run_info['text'])
                    self._apply_run_font(run, run_info.get('font', {}))
            else:
                run = para.add_run(para_info['text'])
                font_info = para_info.get('font', {})
                if font_info:
                    self._apply_run_font(run, font_info)
                else:
                    run.font.size = Pt(12)

        return para

    def _apply_run_font(self, run, font_info: dict):
        """应用文本块字体格式"""
        if font_info.get('name'):
            run.font.name = font_info['name']
        if font_info.get('size'):
            run.font.size = font_info['size']
        if font_info.get('bold'):
            run.font.bold = font_info['bold']
        if font_info.get('italic'):
            run.font.italic = font_info['italic']
        if font_info.get('underline'):
            run.font.underline = font_info['underline']
        if font_info.get('color'):
            run.font.color.rgb = RGBColor.from_string(font_info['color'])

    def _create_table(self, doc: Document, table_data: dict):
        """创建表格"""
        if not table_data.get('data'):
            return None

        rows = len(table_data['data'])
        cols = max([len(row) for row in table_data['data']]) if rows > 0 else 0

        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid' if not table_data.get('style') else table_data['style']

        for i, row_data in enumerate(table_data['data']):
            row = table.rows[i]
            for j, cell_info in enumerate(row_data):
                if isinstance(cell_info, dict):
                    cell = row.cells[j]
                    cell.text = cell_info.get('text', '')

                    for para_info in cell_info.get('paragraphs', []):
                        para = cell.add_paragraph()
                        para.text = para_info.get('text', '')
                else:
                    row.cells[j].text = str(cell_info)

        return table

    def _create_header(self, doc: Document, header_info: dict):
        """创建页眉"""
        section = doc.sections[-1]
        header = section.header

        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = header_info.get('text', '')

        alignment_map = {
            'CENTER': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'LEFT': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'RIGHT': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }

        if header_info.get('alignment') in alignment_map:
            para.alignment = alignment_map[header_info['alignment']]

    def _create_footer(self, doc: Document, footer_info: dict):
        """创建页脚"""
        section = doc.sections[-1]
        footer = section.footer

        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = footer_info.get('text', '')

        alignment_map = {
            'CENTER': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'LEFT': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'RIGHT': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }

        if footer_info.get('alignment') in alignment_map:
            para.alignment = alignment_map[footer_info['alignment']]

    def create_document_to_bytes(self, content: dict) -> bytes:
        """创建文档并返回字节流"""
        doc = Document()

        if content.get('title'):
            title = doc.add_heading(content['title'], level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for para_info in content.get('paragraphs', []):
            para = self._create_paragraph(doc, para_info)

        for table_data in content.get('tables', []):
            if table_data:
                table = self._create_table(doc, table_data)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def extract_footnote_details(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        提取脚注详细内容

        Args:
            docx_path: 文档路径

        Returns:
            list: 脚注列表
        """
        return self._extract_footnotes_from_xml(docx_path)

    def extract_endnote_details(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        提取尾注详细内容

        Args:
            docx_path: 文档路径

        Returns:
            list: 尾注列表
        """
        return self._extract_endnotes_from_xml(docx_path)


def create_doc_processor() -> DocProcessor:
    """
    工厂函数 - 创建文档处理器实例

    Returns:
        DocProcessor: 文档处理器实例
    """
    return DocProcessor()
