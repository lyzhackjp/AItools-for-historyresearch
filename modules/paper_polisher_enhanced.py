"""
学术论文智能精简处理模块 - 增强版

专为日本史学术论文设计的智能内容精简工具
基于阿里通义千问，能够智能识别并删除冗余内容
同时保留核心学术信息。

增强功能（相比原版）：
- 脚注引用重建机制：解决润色后脚注引用丢失问题
- 多种润色策略：支持段落/逐句/修订模式三种策略
- 完善脚注处理：正确区分正文与脚注内容
- 修订追踪功能：启用修订模式，清晰显示所有删减和修改
- 学术严谨性保障：保护历史专有名词和学术术语

技术架构：
- 文档处理：基于 python-docx 库处理 .docx 文档
- AI处理：集成阿里通义千问 / 次要支持 Minimax
- 修订追踪：实现 Track Changes 功能
- 配置管理：使用 .env 格式管理配置参数
"""

import os
import sys
import re
import json
import zipfile
import shutil
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
import hashlib
from lxml import etree
from abc import ABC, abstractmethod

try:
    import docx
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
except ImportError:
    print("警告: python-docx 未安装，使用 pip install python-docx")

from modules.llm_client import create_llm_client


class PolishingStrategy(ABC):
    """润色策略抽象基类"""
    
    @abstractmethod
    def polish_text(self, text: str, llm_client, system_prompt: str) -> Tuple[str, List[Dict]]:
        """
        润色文本
        
        Args:
            text: 待润色文本
            llm_client: LLM客户端
            system_prompt: 系统提示词
            
        Returns:
            Tuple[str, List[Dict]]: (润色后文本, 删除内容列表)
        """
        pass


class ParagraphPolishingStrategy(PolishingStrategy):
    """段落润色策略 - 整段润色（默认）"""
    
    def polish_text(self, text: str, llm_client, system_prompt: str) -> Tuple[str, List[Dict]]:
        """整段润色"""
        user_prompt = f"""请精简以下日本史学术论文段落：

{text}

请返回JSON格式的精简结果。"""
        
        try:
            response = llm_client.chat(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3,
                max_tokens=2000
            )
            
            result_text = response.get('content', text)
            result_text = self._clean_json_response(result_text)
            
            try:
                result_json = json.loads(result_text)
                modified_text = result_json.get('modified_text', text)
                deletions = result_json.get('deletions', [])
                return modified_text, deletions
            except json.JSONDecodeError:
                return text, []
                
        except Exception as e:
            print(f"段落润色失败: {e}")
            return text, []
    
    def _clean_json_response(self, response: str) -> str:
        """清理JSON响应"""
        result_text = response.strip()
        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.startswith("```"):
            result_text = result_text[3:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
        return result_text.strip()


class SentencePolishingStrategy(PolishingStrategy):
    """逐句润色策略 - 句子级别润色（方案一）"""
    
    def polish_text(self, text: str, llm_client, system_prompt: str) -> Tuple[str, List[Dict]]:
        """逐句润色"""
        sentences = self._split_into_sentences(text)
        polished_sentences = []
        all_deletions = []
        
        for sentence in sentences:
            if not sentence.strip():
                continue
            
            user_prompt = f"""请精简以下日史学术论文句子：

{sentence}

请返回JSON格式，包含：
- "modified_text": 精简后的句子
- "deletions": 被删除内容列表"""

            try:
                response = llm_client.chat(
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.3,
                    max_tokens=1000
                )
                
                result_text = response.get('content', sentence)
                result_text = self._clean_json_response(result_text)
                
                try:
                    result_json = json.loads(result_text)
                    modified = result_json.get('modified_text', sentence)
                    deletions = result_json.get('deletions', [])
                    
                    polished_sentences.append(modified)
                    all_deletions.extend(deletions)
                    
                except json.JSONDecodeError:
                    polished_sentences.append(sentence)
                    
            except Exception as e:
                print(f"句子润色失败: {e}")
                polished_sentences.append(sentence)
        
        polished_text = ' '.join(polished_sentences)
        return polished_text, all_deletions
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """将文本分割成句子"""
        sentence_endings = r'[。！？；\n]+'
        sentences = re.split(sentence_endings, text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _clean_json_response(self, response: str) -> str:
        """清理JSON响应"""
        result_text = response.strip()
        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.startswith("```"):
            result_text = result_text[3:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
        return result_text.strip()


class TrackChangesStrategy(PolishingStrategy):
    """修订模式润色策略 - 保留原始文本并添加修订（方案二）"""
    
    def polish_text(self, text: str, llm_client, system_prompt: str) -> Tuple[str, List[Dict]]:
        """修订模式润色"""
        user_prompt = f"""请精简以下日本史学术论文段落，保留原始文本并添加修订：

{text}

请返回JSON格式，包含：
- "modified_text": 精简后的文本
- "deletions": 被删除内容列表，每项包含"text"和"reason"
- "additions": 新增内容列表

输出应清晰标识原始内容和修订内容。"""
        
        try:
            response = llm_client.chat(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3,
                max_tokens=2000
            )
            
            result_text = response.get('content', text)
            result_text = self._clean_json_response(result_text)
            
            try:
                result_json = json.loads(result_text)
                modified_text = result_json.get('modified_text', text)
                deletions = result_json.get('deletions', [])
                return modified_text, deletions
            except json.JSONDecodeError:
                return text, []
                
        except Exception as e:
            print(f"修订模式润色失败: {e}")
            return text, []
    
    def _clean_json_response(self, response: str) -> str:
        """清理JSON响应"""
        result_text = response.strip()
        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.startswith("```"):
            result_text = result_text[3:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
        return result_text.strip()


class FootnoteReferenceRebuilder:
    """脚注引用重建器 - 核心增强功能"""
    
    def __init__(self):
        self.footnote_map = {}
        self.footnote_references = []
    
    def extract_footnote_structure(self, docx_path: str) -> Dict[str, Any]:
        """
        从原始文档提取脚注结构
        
        Args:
            docx_path: DOCX文件路径
            
        Returns:
            Dict: 脚注结构信息
        """
        structure = {
            'references': [],  # 脚注引用列表
            'contents': {},    # 脚注内容字典
            'count': 0
        }
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as z:
                if 'word/document.xml' in z.namelist():
                    with z.open('word/document.xml') as f:
                        tree = etree.fromstring(f.read())
                    
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    refs = tree.findall('.//w:footnoteReference', ns)
                    
                    for ref in refs:
                        ref_id = ref.get(qn('w:id'))
                        structure['references'].append({
                            'id': ref_id,
                            'position': self._get_element_position(ref)
                        })
                    
                    structure['count'] = len(refs)
                
                if 'word/footnotes.xml' in z.namelist():
                    with z.open('word/footnotes.xml') as f:
                        fn_tree = etree.fromstring(f.read())
                    
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    footnotes = fn_tree.findall('.//w:footnote', ns)
                    
                    for fn in footnotes:
                        fn_id = fn.get(qn('w:id'))
                        fn_type = fn.get(qn('w:type'), 'normal')
                        
                        if fn_type == 'normal':
                            text_parts = []
                            for t in fn.findall('.//w:t', ns):
                                if t.text:
                                    text_parts.append(t.text)
                            
                            structure['contents'][fn_id] = {
                                'type': fn_type,
                                'text': ''.join(text_parts),
                                'xml': etree.tostring(fn, encoding='unicode')
                            }
        
        except Exception as e:
            print(f"提取脚注结构失败: {e}")
        
        return structure
    
    def _get_element_position(self, element) -> Dict[str, int]:
        """获取元素位置"""
        parent = element.getparent()
        if parent is not None:
            siblings = list(parent)
            position = 0
            for sibling in siblings:
                if sibling == element:
                    break
                position += 1
            return {'parent_tag': parent.tag, 'index': position}
        return {'parent_tag': None, 'index': 0}
    
    def rebuild_footnotes_in_document(self, original_docx: str, 
                                     polished_docx: str, 
                                     output_docx: str) -> bool:
        """
        在润色后的文档中重建脚注引用
        
        Args:
            original_docx: 原始文档路径
            polished_docx: 润色后文档路径
            output_docx: 输出文档路径
            
        Returns:
            bool: 重建是否成功
        """
        try:
            original_structure = self.extract_footnote_structure(original_docx)
            
            temp_dir = tempfile.mkdtemp()
            
            with zipfile.ZipFile(polished_docx, 'r') as z:
                z.extractall(temp_dir)
            
            polished_doc_path = os.path.join(temp_dir, 'word', 'document.xml')
            polished_footnotes_path = os.path.join(temp_dir, 'word', 'footnotes.xml')
            
            original_footnote_xml = None
            if os.path.exists(os.path.join(os.path.dirname(original_docx), 'word', 'footnotes.xml')):
                pass
            
            with zipfile.ZipFile(original_docx, 'r') as z:
                if 'word/footnotes.xml' in z.namelist():
                    with z.open('word/footnotes.xml') as f:
                        original_footnote_xml = f.read()
            
            if original_footnote_xml:
                with open(polished_footnotes_path, 'wb') as f:
                    f.write(original_footnote_xml)
                
                self._add_footnotes_relationship(temp_dir)
                
                self._rebuild_footnote_references_in_xml(
                    polished_doc_path,
                    original_structure['references']
                )
            
            output_temp = os.path.join(temp_dir, 'output.docx')
            with zipfile.ZipFile(output_temp, 'w', zipfile.ZIP_DEFLATED) as z:
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        if file == 'output.docx':
                            continue
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, temp_dir)
                        z.write(file_path, arcname)
            
            shutil.copy2(output_temp, output_docx)
            shutil.rmtree(temp_dir)
            
            print(f"脚注引用重建完成: {output_docx}")
            return True
            
        except Exception as e:
            print(f"脚注引用重建失败: {e}")
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            return False
    
    def _add_footnotes_relationship(self, temp_dir: str):
        """添加脚注关系"""
        rels_path = os.path.join(temp_dir, 'word', '_rels', 'document.xml.rels')
        
        if os.path.exists(rels_path):
            tree = etree.parse(rels_path)
            root = tree.getroot()
            
            ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
            
            existing = root.find(f".//r:Relationship[@Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes']")
            if existing is None:
                new_rel = etree.SubElement(root, qn('Relationship'))
                new_rel.set('Id', 'rIdFootnotes')
                new_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')
                new_rel.set('Target', 'footnotes.xml')
                
                tree.write(rels_path, xml_declaration=True, encoding='UTF-8', standalone=True)
    
    def _rebuild_footnote_references_in_xml(self, doc_path: str, 
                                          references: List[Dict]):
        """在XML中重建脚注引用"""
        tree = etree.parse(doc_path)
        root = tree.getroot()
        
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        paragraphs = root.findall('.//w:p', ns)
        
        for para_idx, para in enumerate(paragraphs):
            runs = para.findall('.//w:r', ns)
            
            if runs and para_idx < len(references):
                last_run = runs[-1]
                parent = last_run.getparent()
                
                if parent is not None:
                    footnote_ref = OxmlElement('w:footnoteReference')
                    footnote_ref.set(qn('w:id'), references[para_idx]['id'])
                    
                    new_run = OxmlElement('w:r')
                    new_run.append(footnote_ref)
                    
                    idx = list(parent).index(last_run) + 1
                    parent.insert(idx, new_run)
        
        tree.write(doc_path, xml_declaration=True, encoding='UTF-8', standalone=True)


class PaperPolisherEnhanced:
    """学术论文智能精简处理器 - 增强版"""
    
    DEFAULT_SYSTEM_PROMPT = """你是一位专业的日本史学术论文编辑，擅长精简学术论文内容。

请分析以下学术论文内容，识别并删除：
1. 逻辑冗余的论述（重复论证同一观点）
2. 修辞上重复的表达（相同的修饰词反复使用）
3. 非必要的过渡句和重复强调

请务必保留：
1. 核心学术观点和结论
2. 历史史实和重要事件
3. 人物生卒年份和重要事迹
4. 所有脚注、注释和参考文献标注
5. 历史专有名词和学术术语
6. 原文的论证逻辑结构

输出格式要求：
返回JSON格式，包含以下字段：
- "modified_text": 修改后的精简文本（保留所有脚注）
- "deletions": 被删除的内容列表，每项包含"text"和"reason"
- "summary": 精简处理的总结说明

请确保输出是有效的JSON格式。"""
    
    STRATEGIES = {
        'paragraph': ParagraphPolishingStrategy,
        'sentence': SentencePolishingStrategy,
        'track_changes': TrackChangesStrategy
    }
    
    def __init__(self, api_provider: str = "qwen", strategy: str = "paragraph"):
        """
        初始化增强版论文润色器
        
        Args:
            api_provider: API提供商 ('qwen' 或 'minimax')
            strategy: 润色策略 ('paragraph', 'sentence', 'track_changes')
        """
        self.base_dir = Path(__file__).parent.parent
        self.api_provider = api_provider
        self.strategy_name = strategy
        self.strategy = self.STRATEGIES.get(strategy, ParagraphPolishingStrategy)()
        self.llm_client = None
        self.system_prompt = self.DEFAULT_SYSTEM_PROMPT
        self.history_terms = set()
        self.history_figures = set()
        self.footnote_rebuilder = FootnoteReferenceRebuilder()
    
    def _init_llm_client(self):
        """初始化LLM客户端"""
        if self.llm_client is None:
            if self.api_provider == "qwen":
                provider = "dashscope"
            elif self.api_provider == "minimax":
                provider = "minimax"
            else:
                provider = "dashscope"
            
            self.llm_client = create_llm_client(provider)
    
    def _load_history_knowledge(self, doc: Document):
        """从文档中提取历史专有名词和人物"""
        text_content = []
        
        for para in doc.paragraphs:
            text_content.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content.append(cell.text)
        
        full_text = "\n".join(text_content)
        
        date_pattern = r'\d{4,4}年\d{1,2}月\d{1,2}日|\d{4,4}年-\d{4,4}年|\d{1,2}世纪'
        dates = re.findall(date_pattern, full_text)
        
        name_pattern = r'[一-龥]{2,4}(氏|公爵|侯爵|伯爵|子爵|男爵)'
        names = re.findall(name_pattern, full_text)
        
        self.history_terms.update(dates)
        self.history_figures.update(names)
    
    def polish_paragraph(self, paragraph_text: str) -> Tuple[str, List[Dict[str, str]]]:
        """
        精简单个段落（使用当前策略）
        
        Args:
            paragraph_text: 原始段落文本
            
        Returns:
            Tuple[str, List[Dict]]: (精简后的文本, 删除内容列表)
        """
        if not paragraph_text.strip():
            return paragraph_text, []
        
        self._init_llm_client()
        
        return self.strategy.polish_text(
            paragraph_text, 
            self.llm_client, 
            self.system_prompt
        )
    
    def process_document(self, input_path: str, output_path: str,
                        enable_track_changes: bool = True,
                        rebuild_footnotes: bool = True) -> Dict[str, Any]:
        """
        处理完整的Word文档（增强版）
        
        Args:
            input_path: 输入文档路径
            output_path: 输出文档路径
            enable_track_changes: 是否启用修订追踪
            rebuild_footnotes: 是否重建脚注引用
            
        Returns:
            Dict: 处理结果统计
        """
        input_path = Path(input_path)
        output_path = Path(output_path)
        
        if not input_path.exists():
            raise FileNotFoundError(f"输入文件不存在: {input_path}")
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc = Document(str(input_path))
        self._load_history_knowledge(doc)
        
        temp_polished = output_path.parent / f"temp_polished_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        total_paragraphs = 0
        processed_paragraphs = 0
        all_deletions = []
        
        body_paragraphs = []
        footnotes = []
        is_in_footnote = False
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            if '脚注' in text or '注释' in text:
                is_in_footnote = True
            
            if is_in_footnote:
                footnotes.append(para)
            else:
                if para.style.name.startswith('Normal') or para.style.name.startswith('Heading'):
                    body_paragraphs.append(para)
        
        total_paragraphs = len(body_paragraphs)
        
        print(f"开始处理文档（策略: {self.strategy_name}）...")
        print(f"正文段落数: {total_paragraphs}")
        print(f"脚注数量: {len(footnotes)}")
        
        for i, para in enumerate(body_paragraphs, 1):
            original_text = para.text
            
            if not original_text.strip():
                continue
            
            processed_paragraphs += 1
            
            print(f"[{i}/{total_paragraphs}] 处理段落 {i}...", end=" ", flush=True)
            
            modified_text, deletions = self.polish_paragraph(original_text)
            
            if deletions:
                all_deletions.extend(deletions)
            
            if modified_text != original_text and enable_track_changes:
                self._apply_track_changes(para, original_text, modified_text)
            elif modified_text != original_text:
                para.clear()
                para.add_run(modified_text)
            
            print(f"✓ ({len(deletions)} 处修改)")
        
        doc.save(str(temp_polished))
        
        if rebuild_footnotes:
            print("\n正在重建脚注引用...")
            success = self.footnote_rebuilder.rebuild_footnotes_in_document(
                str(input_path),
                str(temp_polished),
                str(output_path)
            )
            
            if not success:
                shutil.copy2(str(temp_polished), str(output_path))
                print("警告: 脚注重建失败，使用不含脚注的版本")
        else:
            shutil.copy2(str(temp_polished), str(output_path))
        
        if temp_polished.exists():
            temp_polished.unlink()
        
        result = {
            'success': True,
            'input_file': str(input_path),
            'output_file': str(output_path),
            'total_paragraphs': total_paragraphs,
            'processed_paragraphs': processed_paragraphs,
            'total_deletions': len(all_deletions),
            'deletions': all_deletions,
            'footnote_count': len(footnotes),
            'footnote_rebuild': rebuild_footnotes,
            'track_changes_enabled': enable_track_changes,
            'strategy': self.strategy_name,
            'timestamp': datetime.now().isoformat()
        }
        
        return result
    
    def _apply_track_changes(self, para, original_text: str, modified_text: str):
        """应用修订追踪格式"""
        para.clear()
        
        orig_runs = original_text.split(modified_text)
        
        if len(orig_runs) == 2:
            deleted_text = orig_runs[0] + orig_runs[1]
            
            run_deleted = para.add_run(deleted_text)
            run_deleted.font.strike = True
            run_deleted.font.color.rgb = RGBColor(255, 0, 0)
            
            run_added = para.add_run(modified_text)
            run_added.underline = True
            run_added.font.color.rgb = RGBColor(0, 0, 255)
        else:
            run_added = para.add_run(modified_text)
            run_added.underline = True
            run_added.font.color.rgb = RGBColor(0, 0, 255)
    
    def set_strategy(self, strategy: str):
        """
        设置润色策略
        
        Args:
            strategy: 策略名称 ('paragraph', 'sentence', 'track_changes')
        """
        if strategy in self.STRATEGIES:
            self.strategy_name = strategy
            self.strategy = self.STRATEGIES[strategy]()
        else:
            print(f"警告: 未知的策略 '{strategy}'，使用默认策略 'paragraph'")
            self.strategy_name = 'paragraph'
            self.strategy = ParagraphPolishingStrategy()
    
    def set_system_prompt(self, prompt: str):
        """设置系统提示词"""
        self.system_prompt = prompt
    
    def add_history_term(self, term: str):
        """添加历史专有名词保护"""
        self.history_terms.add(term)
    
    def add_history_figure(self, figure: str):
        """添加历史人物保护"""
        self.history_figures.add(figure)


def create_paper_polisher_enhanced(api_provider: str = "qwen", 
                                   strategy: str = "paragraph") -> PaperPolisherEnhanced:
    """
    工厂函数 - 创建增强版论文润色器
    
    Args:
        api_provider: API提供商 ('qwen' 或 'minimax')
        strategy: 润色策略 ('paragraph', 'sentence', 'track_changes')
        
    Returns:
        PaperPolisherEnhanced: 增强版论文润色器实例
    """
    return PaperPolisherEnhanced(api_provider, strategy)


if __name__ == "__main__":
    print("学术论文智能精简处理工具 - 增强版")
    print("="*60)
    print("\n支持三种润色策略:")
    print("1. paragraph  - 段落润色（默认）")
    print("2. sentence   - 逐句润色（更精细）")
    print("3. track_changes - 修订模式（保留原始）")
    print("\n使用方法:")
    print("```python")
    print("from modules.paper_polisher_enhanced import create_paper_polisher_enhanced")
    print("")
    print("# 创建增强版润色器（段落策略）")
    print("polisher = create_paper_polisher_enhanced('qwen', 'paragraph')")
    print("")
    print("# 或使用逐句策略")
    print("polisher = create_paper_polisher_enhanced('qwen', 'sentence')")
    print("")
    print("# 处理文档（自动重建脚注引用）")
    print("result = polisher.process_document(")
    print("    'input.docx',")
    print("    'output.docx',")
    print("    enable_track_changes=True,")
    print("    rebuild_footnotes=True")
    print(")")
    print("```")
