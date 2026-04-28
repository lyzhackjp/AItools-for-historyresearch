import sys, os
sys.path.insert(0, r'C:\Users\lyzha\Desktop\AItools-for-historyresearch')
os.chdir(r'C:\Users\lyzha\Desktop\AItools-for-historyresearch')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read the paper draft
draft_path = "./workflow_output/paper_draft.md"
final_path = "./workflow_output/final_paper.md"
output_path = "./workflow_output/paper_draft.docx"

# Use final paper if available, otherwise draft
if os.path.exists(final_path):
    paper_path = final_path
    print(f"Using final paper: {final_path}")
else:
    paper_path = draft_path
    print(f"Using draft: {draft_path}")

with open(paper_path, "r", encoding="utf-8") as f:
    content = f.read()

print(f"Paper content: {len(content)} chars")

# Create Word document
doc = Document()
doc.core_properties.title = "Tudor England - History Research Paper"
doc.core_properties.author = "AI History Research Assistant"

# Title
title = doc.add_heading("Tudor England", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph("A History Research Paper")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # spacing

# Add content paragraph by paragraph
# Split by double newlines for paragraphs
sections = content.split("\n\n")
for section in sections:
    section = section.strip()
    if not section:
        continue
    # Check if it's a heading
    if section.startswith("#"):
        # Markdown heading
        level = section.count("#")
        heading_text = section.lstrip("#").strip()
        if level == 1:
            doc.add_heading(heading_text, level=1)
        elif level == 2:
            doc.add_heading(heading_text, level=2)
        else:
            doc.add_heading(heading_text, level=3)
    else:
        p = doc.add_paragraph(section)
        p.paragraph_format.space_after = Pt(12)

doc.save(output_path)
print(f"Word document saved to: {output_path}")
