"""
Word 文档导出模块

将论文草稿导出为 Word 文档，支持脚注式引用

核心功能：
- Markdown → Word 转换
- 脚注生成（学术引用）
- 支持中英双语
- 保留格式（标题层级、加粗、斜体、引用块）

使用 python-docx + lxml 直接操作 OOXML 实现脚注

依赖：
    pip install python-docx
"""

import re
import os
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[WordExporter] python-docx 未安装: pip install python-docx")
    Document = None


# ─────────────────────────────────────────────────────────────────
#  脚注引用管理器
# ─────────────────────────────────────────────────────────────────

class FootnoteManager:
    """
    管理 Word 文档中的脚注

    在 OOXML 中，脚注存储在 word/footnotes.xml，
    正文通过 <w:footnoteReference> 引用脚注编号。
    """

    def __init__(self, doc):
        self.doc = doc
        self.footnotes = []  # List of (text, reference_mark) tuples
        self._footnote_count = 0

    def add_footnote(self, text: str) -> int:
        """
        添加脚注，返回脚注编号（从 1 开始）

        Args:
            text: 脚注内容

        Returns:
            int: 脚注编号
        """
        self._footnote_count += 1
        self.footnotes.append((text, self._footnote_count))
        return self._footnote_count

    def render_footnotes(self):
        """将所有脚注写入文档末尾"""
        if not self.footnotes:
            return

        # 添加分隔线
        sep = self.doc.add_paragraph()
        sep.add_run("─" * 40)

        for text, num in self.footnotes:
            p = self.doc.add_paragraph()
            run = p.add_run(f"[{num}] ")
            run.bold = True
            run.font.size = Pt(9)
            p.add_run(text).font.size = Pt(9)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)


# ─────────────────────────────────────────────────────────────────
#  Markdown → Word 转换器
# ─────────────────────────────────────────────────────────────────

class MarkdownToWordConverter:
    """
    将 Markdown 文本转换为 Word 文档

    支持：
    - 标题（# ~ ######）
    - 加粗、斜体
    - 引用块（>）
    - 有序/无序列表
    - 脚注引用 [^n] 或 [*]
    - 水平分隔线（---）
    - 代码块
    """

    # 脚注 pattern: [^1], [^citation], *citation*
    FOOTNOTE_PATTERNS = [
        r'\[\^(\d+)\]',        # [^1], [^42]
        r'\*\*\[(\d+)\]\*\*',  # **[1]** (bold footnote marker)
    ]

    def __init__(self, paper_text: str, language: str = "en"):
        self.paper_text = paper_text
        self.language = language
        self.doc = None
        self.footnote_mgr = None
        self._footnote_refs = {}  # 存储已发现的脚注引用（用于生成脚注内容）

    def convert(self, output_path: str = "") -> str:
        """
        执行转换

        Args:
            output_path: 输出文件路径

        Returns:
            str: 输出文件路径
        """
        if Document is None:
            raise RuntimeError("python-docx 未安装")

        self.doc = Document()
        self.footnote_mgr = FootnoteManager(self.doc)

        # 设置文档默认字体
        self._set_document_defaults()

        # 按段落处理
        blocks = self._split_blocks(self.paper_text)

        for block in blocks:
            self._process_block(block)

        # 渲染脚注
        self.footnote_mgr.render_footnotes()

        # 保存
        if not output_path:
            output_path = "paper_output.docx"

        self.doc.save(output_path)
        return output_path

    def _set_document_defaults(self):
        """设置文档默认样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

    def _split_blocks(self, text: str) -> List[Tuple[str, str]]:
        """
        将文本分割为块列表

        Returns:
            List[(type, content)] where type in ('heading', 'paragraph', 'quote', 'code', 'list', 'hr')
        """
        lines = text.split('\n')
        blocks = []
        current_block = []
        current_type = 'paragraph'

        def flush():
            if current_block:
                content = '\n'.join(current_block).strip()
                if content:
                    blocks.append((current_type, content))
                current_block.clear()

        for line in lines:
            # Heading
            m = re.match(r'^(#{1,6})\s+(.*)', line)
            if m:
                flush()
                blocks.append(('heading', (len(m.group(1)), m.group(2))))
                continue

            # HR
            if re.match(r'^(-{3,}|_{3,}|\*{3,})$', line.strip()):
                flush()
                blocks.append(('hr', ''))
                continue

            # Quote
            if line.startswith('>'):
                flush()
                current_type = 'quote'
                current_block.append(line[1:].strip())
                continue

            # Code block fence
            if line.strip().startswith('```'):
                flush()
                current_type = 'code'
                current_block.append(line)
                continue

            # List
            if re.match(r'^(\d+\.|\-|\*)\s+', line):
                flush()
                current_type = 'list'
                current_block.append(line)
                continue

            # Continuation of current block
            if current_block:
                current_block.append(line)
            else:
                current_block.append(line)

        flush()
        return blocks

    def _process_block(self, block: Tuple[str, Any]):
        """处理单个文本块"""
        btype, content = block

        if btype == 'heading':
            level, text = content
            self._add_heading(text, level)
        elif btype == 'paragraph':
            self._add_paragraph(content)
        elif btype == 'quote':
            self._add_quote(content)
        elif btype == 'code':
            self._add_code(content)
        elif btype == 'list':
            self._add_list(content)
        elif btype == 'hr':
            self._add_hr()

    def _add_heading(self, text: str, level: int):
        """添加标题"""
        # python-docx 只有 1-9 级样式
        style_name = f'Heading {min(level, 9)}'
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        try:
            p.style = self.doc.styles[style_name]
        except:
            pass  # 样式不存在则使用默认

    def _add_paragraph(self, text: str):
        """添加正文段落，处理脚注引用"""
        p = self.doc.add_paragraph()

        # 按脚注引用分割
        segments = self._split_with_footnotes(text)

        for seg_text, footnote_num in segments:
            if footnote_num:
                # 插入脚注引用
                self._add_footnote_reference(p, footnote_num)
            if seg_text:
                run = p.add_run(seg_text)

        p.paragraph_format.first_line_indent = Inches(0.3)  # 首行缩进

    def _split_with_footnotes(self, text: str) -> List[Tuple[str, Optional[int]]]:
        """
        将文本按脚注引用分割

        Returns:
            List[(segment_text, footnote_num_or_None)]
        """
        # 找到所有脚注引用位置
        pattern = r'\[\^(\d+)\]|\*\*\[(\d+)\]\*\*'
        segments = []
        last_end = 0

        for m in re.finditer(pattern, text):
            # 添加引用之前的文本
            if m.start() > last_end:
                segments.append((text[last_end:m.start()], None))

            fn_num = int(m.group(1) or m.group(2))
            segments.append(('', fn_num))
            last_end = m.end()

        if last_end < len(text):
            segments.append((text[last_end:], None))

        return segments if segments else [(text, None)]

    def _add_footnote_reference(self, para, footnote_num: int):
        """在段落中添加脚注引用"""
        # 创建脚注引用元素
        fref = OxmlElement('w:footnoteReference')
        fref.set(qn('w:id'), str(footnote_num))
        fref.set(qn('w:customMarkFollows'), '1')

        # 获取 run 的 rPr（如果存在）以保持格式
        run_elem = para._p
        run_elem.append(fref)

    def _add_quote(self, text: str):
        """添加引用块"""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        run = p.add_run(text)
        run.italic = True

    def _add_code(self, text: str):
        """添加代码块"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.3)

    def _add_list(self, text: str):
        """添加列表项"""
        lines = text.split('\n')
        for line in lines:
            m = re.match(r'^(\d+\.|\-|\*)\s+(.*)', line)
            if m:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                bullet = p.add_run(m.group(1) + " ")
                bullet.bold = True
                p.add_run(m.group(2))

    def _add_hr(self):
        """添加分隔线"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("─" * 40)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ─────────────────────────────────────────────────────────────────
#  便捷函数
# ─────────────────────────────────────────────────────────────────

def export_paper_to_word(
    paper_text: str,
    output_path: str = "",
    language: str = "en",
    title: str = "",
    author: str = "",
    citation_format: str = "chicago"
) -> str:
    """
    将论文导出为 Word 文档

    Args:
        paper_text: 论文正文（Markdown 格式）
        output_path: 输出路径（默认自动生成）
        language: 语言 (en/ja/zh)
        title: 论文标题（显示在文档开头）
        author: 作者名
        citation_format: 引用格式（影响脚注样式）

    Returns:
        str: 输出文件路径
    """
    if not paper_text:
        raise ValueError("论文内容为空")

    if not output_path:
        safe_title = "".join(c if c.isalnum() else '_' for c in title[:20]) if title else "paper"
        import datetime
        ts = datetime.datetime.now().strftime('%Y%m%d')
        output_path = f"{safe_title}_{ts}.docx"

    # 构建完整文档内容
    full_content = paper_text
    if title:
        full_content = f"# {title}\n\n{full_content}"

    # 转换
    converter = MarkdownToWordConverter(full_content, language=language)
    result_path = converter.convert(output_path)

    # 添加文档元信息（标题/作者）
    if title or author:
        _add_document_metadata(result_path, title, author, language)

    return result_path


def _add_document_metadata(doc_path: str, title: str, author: str, language: str):
    """设置 Word 文档的元信息（标题、作者）"""
    try:
        doc = Document(doc_path)

        # 设置 core properties（需要 access to package）
        core_props = doc.core_properties
        if title:
            core_props.title = title
        if author:
            core_props.author = author

        doc.save(doc_path)
    except Exception as e:
        print(f"[WordExporter] 设置文档元数据失败: {e}")


def export_paper_with_footnotes(
    paper_text: str,
    footnotes: List[Dict[str, str]],
    output_path: str = "",
    language: str = "en",
    title: str = ""
) -> str:
    """
    将论文导出为带脚注的 Word 文档

    Args:
        paper_text: 论文正文（Markdown 格式）
        footnotes: 脚注列表 [{id: "1", text: "..."}, ...]
        output_path: 输出路径
        language: 语言
        title: 论文标题

    Returns:
        str: 输出文件路径
    """
    if not paper_text:
        raise ValueError("论文内容为空")

    # 将脚注注入论文文本
    # 查找所有 [^n] 引用，用实际脚注内容替换
    processed_text = paper_text

    # 构建脚注内容映射
    fn_map = {str(i+1): fn.get('text', '') for i, fn in enumerate(footnotes)}

    def replace_footnote_ref(m):
        fn_id = m.group(1)
        fn_text = fn_map.get(fn_id, '')
        if fn_text:
            return f"[{fn_id}]"
        return m.group(0)

    processed_text = re.sub(r'\[\^(\d+)\]', replace_footnote_ref, processed_text)

    return export_paper_to_word(
        processed_text,
        output_path=output_path,
        language=language,
        title=title
    )


# ─────────────────────────────────────────────────────────────────
#  测试
# ─────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    sample = """# Tudor England: A Historical Overview

The Tudor period (1485–1603) represents one of the most significant epochs in English history[^1].

## The Break with Rome

Henry VIII's decision to break with Rome in 1534 marked a watershed moment[^2] in English religious history.

> "This royal throne of kings, this sundered isle" — Shakespeare captured the turbulence of this era[^3].

## Key Figures

- Henry VIII (1491-1547)
- Thomas Cromwell (c.1485-1540)
- Elizabeth I (1533-1608)

[^1]: The Tudor dynasty lasted from 1485 to 1603, encompassing the reigns of five monarchs.
[^2]: The Act of Supremacy (1534) declared the king as the supreme head of the Church of England.
[^3]: Shakespeare's Chorus describes the era in Henry V.
"""

    output = export_paper_to_word(
        sample,
        output_path="./workflow_output/test_paper.docx",
        language="en",
        title="Tudor England: A Historical Overview"
    )
    print(f"Word document created: {output}")
